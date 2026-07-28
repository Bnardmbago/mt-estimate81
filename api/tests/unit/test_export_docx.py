from io import BytesIO

from docx import Document

from app.exports.docx import (
    generate_quotation_docx,
    generate_quotation_formal_docx,
    generate_report_docx,
)
from app.proposals.svg_raster import svg_to_png_bytes
from tests.unit.export_fixtures import (
    sample_estimate_with_discount,
    sample_formal_quotation_context,
    sample_quotation_context,
    sample_report_context,
)


def _cell_text(cell) -> str:
    parts: list[str] = []
    for paragraph in cell.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in cell.tables:
        parts.append(_table_text(table))
    return "\n".join(parts)


def _table_text(table) -> str:
    parts: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            text = _cell_text(cell)
            if text.strip():
                parts.append(text)
    return "\n".join(parts)


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        parts.append(_table_text(table))
    return "\n".join(parts)


def test_report_docx_is_valid_docx_zip():
    content = generate_report_docx(sample_report_context())
    assert content[:2] == b"PK"


def test_report_docx_rasterizes_accent_before_cover_text(monkeypatch):
    accent_svg = (
        '<svg xmlns="http://www.w3.org/2000/svg" width="355.6mm" height="215.9mm">'
        '<circle cx="80" cy="80" r="50" fill="#2563eb" fill-opacity=".5"/></svg>'
    )
    png = svg_to_png_bytes(accent_svg)
    assert png is not None
    calls: list[tuple[str, float]] = []

    def rasterize(svg: str, *, scale: float = 1.5):
        calls.append((svg, scale))
        return png

    monkeypatch.setattr("app.exports.docx.svg_to_png_bytes", rasterize, raising=False)
    ctx = sample_report_context(include_cover=True)
    ctx["page"] = {"size": "Legal", "orientation": "landscape"}
    ctx["cover"] = {
        "fields": [
            {
                "key": "title",
                "label": "Title",
                "value": "Estimate Cover",
                "emphasis": "title",
            }
        ],
        "assets": [{"region": "logo", "url": "data:image/png;base64,unused"}],
        "accent_svg": accent_svg,
        "warnings": [],
    }

    content = generate_report_docx(ctx)

    document = Document(BytesIO(content))
    assert calls == [(accent_svg, 1.5)]
    assert len(document.inline_shapes) == 1
    assert document.inline_shapes[0].width > document.inline_shapes[0].height
    assert document.sections[0].page_width.mm == 355.6
    assert document.sections[0].page_height.mm == 215.9
    assert document._element.xml.index("<w:drawing>") < document._element.xml.index(
        "Estimate Cover"
    )
    assert "Estimate Cover" in _docx_text(content)
    assert ctx["cover"]["assets"][0]["region"] == "logo"


def test_report_docx_omits_failed_accent_and_adds_fidelity_warning(monkeypatch):
    monkeypatch.setattr(
        "app.exports.docx.svg_to_png_bytes",
        lambda _svg, **_kwargs: None,
        raising=False,
    )
    ctx = sample_report_context(include_cover=True)
    ctx["cover"] = {
        "fields": [
            {
                "key": "title",
                "label": "Title",
                "value": "Fallback Estimate",
                "emphasis": "title",
            }
        ],
        "assets": [],
        "accent_svg": "<svg xmlns='http://www.w3.org/2000/svg'/>",
        "warnings": [],
    }

    content = generate_report_docx(ctx)

    document = Document(BytesIO(content))
    assert content[:2] == b"PK"
    assert len(document.inline_shapes) == 0
    assert "Fallback Estimate" in _docx_text(content)
    assert (
        "DOCX uses flow-based Cover layout; exact positioning may differ from PDF."
        in ctx["cover"]["warnings"]
    )


def test_report_docx_contains_project_and_cost_sections():
    content = generate_report_docx(sample_report_context())
    text = _docx_text(content)
    assert "Portal Redesign" in text
    assert "ACME Corp" in text
    assert "Development cost" in text
    assert "Development period" in text
    assert "NRC Breakdown (Detailed)" in text
    assert "RC Breakdown (Detailed)" in text


def test_report_docx_excludes_removed_sections():
    content = generate_report_docx(sample_report_context())
    text = _docx_text(content)
    assert "Rate Card Reference" not in text
    assert "AI Confidence" not in text
    assert "Cost Drivers" not in text
    assert "Risks & Gaps" not in text


def test_report_docx_ja_locale():
    content = generate_report_docx(
        sample_report_context(locale="ja", export_user_display_name="山田太郎")
    )
    text = _docx_text(content)
    assert "開発コストの概要" in text
    assert "開発費用" in text
    assert "保守運用費用　月額 /年間" in text
    assert "開発期間" in text
    assert "見積作成者" in text
    assert "山田太郎" in text
    assert "初年度合計コスト" not in text


def test_quotation_docx_contains_title_client_and_totals():
    content = generate_quotation_docx(sample_quotation_context())
    text = _docx_text(content)
    assert "QUOTATION" in text
    assert "ACME Corp" in text
    assert "Development" in text
    assert "Infrastructure Setup" in text
    assert "¥770,000" in text


def test_quotation_docx_contains_quote_number_when_populated():
    content = generate_quotation_docx(
        sample_formal_quotation_context(
            quotation_number="BAI-20260629-001",
            registration_number="T9010001234562",
            contact_person="Tanaka Taro",
        )
    )
    text = _docx_text(content)
    assert "Quotation No." in text
    assert "BAI-20260629-001" in text
    assert "T9010001234562" in text
    assert "Tanaka Taro" in text


def test_quotation_docx_ja_locale():
    content = generate_quotation_docx(sample_quotation_context(locale="ja"))
    text = _docx_text(content)
    assert "見積書" in text
    assert "御中" in text
    assert "開発" in text
    assert "インフラセットアップ" in text


def test_formal_quotation_docx_contains_populated_numbers():
    content = generate_quotation_formal_docx(
        sample_formal_quotation_context(estimate=sample_estimate_with_discount())
    )
    text = _docx_text(content)
    assert "BAI-20260629-001" in text
    assert "T9010001234562" in text
    assert "Development" in text
    assert "Special Discount" in text
    assert "[Notes]" in text
    assert "Campaign Terms" not in text
    assert "*Special Notes" not in text


def test_formal_quotation_docx_ja_discount_row():
    content = generate_quotation_formal_docx(
        sample_formal_quotation_context(
            estimate=sample_estimate_with_discount(),
            locale="ja",
        )
    )
    text = _docx_text(content)
    assert "開発" in text
    assert "特別割引" in text
    assert "【備考】" in text


def test_report_docx_includes_discount_pricing_when_present():
    content = generate_report_docx(
        sample_report_context(estimate=sample_estimate_with_discount())
    )
    text = _docx_text(content)
    assert "Development Cost" in text
    assert "Limited-Time Discount" in text
    assert "Special Price" in text
    assert "Campaign Terms" in text
    assert "¥1,000,000" in text
    assert "30% OFF" in text


def test_report_docx_rc_breakdown_includes_monthly_and_annual_totals():
    content = generate_report_docx(sample_report_context())
    text = _docx_text(content)
    assert "Monthly RC Total" in text
    assert "Annual RC Total" in text
    assert "¥170,000" in text
    assert "¥2,040,000" in text
    assert "Cloud Infrastructure" in text
    assert "Server & database usage" in text
    assert "Maintenance and Support" in text


def test_quotation_docx_includes_discount_as_line_item_when_present():
    content = generate_quotation_docx(
        sample_quotation_context(estimate=sample_estimate_with_discount(), locale="en")
    )
    text = _docx_text(content)
    assert "Development" in text
    assert "Special Discount" in text
    assert "[Notes]" in text
    assert "Limited-Time Discount" not in text
    assert "Campaign Terms" not in text
