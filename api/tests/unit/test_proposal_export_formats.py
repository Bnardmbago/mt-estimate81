"""Unit tests for proposal DOCX/XLSX content parity with the PDF pack."""

from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from openpyxl import load_workbook

from app.presentation.accent_shapes import normalize_accent_shapes
from app.presentation.resolver import PresentationBundle
from app.proposals.export_context import build_proposal_export_context
from app.proposals.export_formats import (
    generate_proposal_docx,
    generate_proposal_markdown,
    generate_proposal_pdf,
    generate_proposal_xlsx,
    render_proposal_html,
)
from app.proposals.export_service import _embed_cover_asset_data
from app.proposals.svg_raster import svg_to_png_bytes


def _sample_ctx(**overrides):
    ctx = {
        "labels": {
            "title": "Project Proposal Pack",
            "toc": "Table of Contents",
            "assessment": "Project Assessment",
            "proposal": "Project Proposal",
            "poc": "Proof of Concept",
            "one_time": "One-time project cost",
            "monthly": "Monthly recurring cost",
            "first_year": "First-year total",
            "timeline": "Project timeline",
            "milestones": "Milestones",
            "project_brief": "Project brief",
            "brief_project_name": "Project name",
            "brief_description": "Project description",
            "brief_business_problem": "Business problem",
            "brief_target_users": "Target users",
            "brief_technology_stack": "Technology stack",
            "brief_constraints": "Constraints",
            "poc_tables": "Tables",
            "poc_diagrams": "Illustrations",
            "poc_milestones": "Proof of Concept milestones",
            "official_poc_cost": "Official Proof of Concept cost",
            "suggested_window": "Suggested validation window",
            "effort_hours": "Estimated effort (hours)",
        },
        "locale": "en",
        "variant": "full",
        "project_name": "Portal Redesign",
        "client_name": "ACME Corp",
        "toc": [
            {"id": "assessment", "title": "Project Assessment", "level": "1"},
            {"id": "proposal", "title": "Project Proposal", "level": "1"},
        ],
        "assessment": {
            "sections": [
                {
                    "id": "overview",
                    "title": "Assessment Overview",
                    "body": "Assessment body text",
                    "bullets": ["Risk A"],
                    "rating": "high",
                }
            ]
        },
        "proposal_body": {
            "sections": [
                {
                    "id": "approach",
                    "title": "Delivery Approach",
                    "body": "We will deliver in phases.",
                    "bullets": ["Phase 1", "Phase 2"],
                }
            ]
        },
        "poc": {
            "project_brief": {
                "project_name": "Portal Redesign",
                "project_description": "Modernize the portal",
                "business_problem": "Legacy UX",
                "target_users": "Employees",
                "technology_stack": "Next.js",
                "constraints": "8 weeks",
            },
            "sections": [
                {
                    "id": "scope",
                    "title": "PoC Scope",
                    "body": "Validate authentication",
                    "bullets": ["SSO"],
                }
            ],
            "tables": [
                {
                    "title": "Effort table",
                    "headers": ["Role", "Hours"],
                    "rows": [["Developer", "40"]],
                }
            ],
            "diagrams": [
                {
                    "title": "PoC Flow",
                    "source": "flowchart LR\n  A --> B",
                }
            ],
            "milestones": [{"name": "PoC kickoff", "date": "2026-09-01"}],
            "official": {
                "total_effort_hours": 80,
                "estimated_one_time_cost_jpy": 800000,
            },
            "suggested_validation_window": "4–6 weeks",
        },
        "diagrams": [
            {
                "title": "Architecture Overview",
                "source": "graph TD\n  Client --> API",
            }
        ],
        "milestones": [{"name": "Go-live", "date": "2026-12-01"}],
        "cost_summary": {
            "one_time_project_cost_jpy": 5000000,
            "monthly_recurring_cost_jpy": 200000,
            "first_year_total_jpy": 7400000,
        },
        "gantt": {
            "project_start_date": "2026-08-01",
            "tasks": [
                {
                    "name": "Design UI",
                    "phase": "design",
                    "start_date": "2026-08-01",
                    "end_date": "2026-08-10",
                    "duration_working_days": 7,
                    "hours": 40,
                }
            ],
            "phases": [],
        },
        "gantt_svg": (
            '<svg xmlns="http://www.w3.org/2000/svg" width="100" height="40">'
            '<rect width="100" height="40" fill="#1E3A5F"/>'
            "</svg>"
        ),
        "include_poc": True,
        "theme": {
            "primary": "1E3A5F",
            "primary_light": "E8EEF4",
            "accent": "2563EB",
            "text_on_primary": "FFFFFF",
            "surface": "F8FAFC",
            "border": "E2E8F0",
            "text_body": "1E293B",
            "text_muted": "64748B",
        },
        "style": {
            "base_font_size_pt": 10,
            "line_spacing": 1.4,
            "margins": {"top_mm": 18, "right_mm": 16, "bottom_mm": 18, "left_mm": 16},
        },
        "layout": {
            "layout": "linear",
            "cover": False,
            "section_chrome": "ruled",
            "columns": 1,
        },
        "include_cover": False,
        "page": {"size": "A4", "orientation": "portrait", "css_size": "A4 portrait"},
        "cover": {
            "fields": [],
            "assets": [],
            "design": {},
            "warnings": [],
        },
        "presentation": {
            "theme_id": "corporate-navy",
            "style_id": "comfortable",
            "template_id": "classic-linear",
            "layout_class": "proposal-layout-linear",
        },
    }
    ctx.update(overrides)
    return ctx


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _accent_shape(
    shape_type: str = "rectangle",
    *,
    shape_id: str = "shape-1",
    z_index: int = 1,
    visible: bool = True,
):
    raw = {
        "id": shape_id,
        "name": shape_type.title(),
        "type": shape_type,
        "visible": visible,
        "geometry": {
            "x_pct": 10,
            "y_pct": 10,
            "width_pct": 30,
            "height_pct": 20,
            "rotation_deg": 0,
            "z_index": z_index,
        },
        "fill": {"mode": "theme", "opacity": 1},
    }
    if shape_type == "polygon":
        raw["points"] = [
            {"x_pct": 0, "y_pct": 0},
            {"x_pct": 100, "y_pct": 0},
            {"x_pct": 50, "y_pct": 100},
        ]
    shapes, warnings = normalize_accent_shapes([raw])
    assert not warnings
    return shapes[0]


def _proposal_for_export():
    return SimpleNamespace(
        locale="en",
        source_snapshot={"project_name": "Accent proposal", "client_name": "ACME"},
        include_poc=False,
        poc=None,
        assessment=None,
        proposal_body=None,
        diagrams=[],
        milestones=[],
        cover_values={},
    )


def test_proposal_html_uses_accent_for_approved_document_highlights():
    base_theme = _sample_ctx()["theme"]
    ctx = _sample_ctx(
        theme={
            **base_theme,
            "accent": "C026D3",
            "callout": "F59E0B",
            "table_highlight": "10B981",
        }
    )

    html = render_proposal_html(ctx)

    assert "--accent: #C026D3" in html
    assert "--callout: #F59E0B" in html
    assert "--table-highlight: #10B981" in html
    assert "border-bottom: 2.5pt solid var(--accent)" in html
    assert "border-left: 4pt solid var(--callout)" in html
    assert "border-bottom: 2pt solid var(--table-highlight)" in html
    assert "border-left: 4pt solid var(--accent)" in html


def test_proposal_docx_nonempty_and_contains_pack_content():
    content = generate_proposal_docx(_sample_ctx())
    assert content[:2] == b"PK"
    assert len(content) > 1000
    text = _docx_text(content)
    assert "Portal Redesign" in text
    assert "Assessment Overview" in text
    assert "Delivery Approach" in text
    assert "PoC Scope" in text
    assert "Architecture Overview" in text
    assert "graph TD" in text
    assert "Client --> API" in text
    assert "Effort table" in text
    assert "Developer" in text
    assert "Go-live" in text
    assert "Modernize the portal" in text


def test_proposal_docx_embeds_gantt_picture_when_svg_present():
    # Prefer real cairosvg rasterization when available in the test environment.
    png = svg_to_png_bytes(_sample_ctx()["gantt_svg"])
    if png is None:
        # Environment without Cairo — skip visual embed assertion.
        return
    content = generate_proposal_docx(_sample_ctx())
    document = Document(BytesIO(content))
    assert document.inline_shapes
    text = _docx_text(content)
    assert "Project timeline" in text


def test_proposal_docx_skips_gantt_picture_when_raster_fails(monkeypatch):
    monkeypatch.setattr(
        "app.proposals.export_formats.svg_to_png_bytes",
        lambda _svg, **_kwargs: None,
    )
    content = generate_proposal_docx(_sample_ctx())
    document = Document(BytesIO(content))
    assert len(document.inline_shapes) == 0


def test_proposal_xlsx_sheets_and_content():
    content = generate_proposal_xlsx(_sample_ctx())
    assert content[:2] == b"PK"
    wb = load_workbook(BytesIO(content))
    assert "Summary" in wb.sheetnames
    assert "Assessment" in wb.sheetnames
    assert "Proposal" in wb.sheetnames
    assert "PoC" in wb.sheetnames
    assert "PoC Tables" in wb.sheetnames
    assert "Diagrams" in wb.sheetnames
    assert "Timeline" in wb.sheetnames
    assert "Milestones" in wb.sheetnames

    summary = wb["Summary"]
    assert summary["B1"].value == "Value" or summary["A2"].value == "Project"
    assert any(cell.value == "Portal Redesign" for row in summary.iter_rows() for cell in row)

    assessment = wb["Assessment"]
    assert any(cell.value == "Assessment Overview" for row in assessment.iter_rows() for cell in row)

    timeline = wb["Timeline"]
    assert any(cell.value == "Design UI" for row in timeline.iter_rows() for cell in row)

    diagrams = wb["Diagrams"]
    values = [cell.value for row in diagrams.iter_rows() for cell in row]
    assert "Architecture Overview" in values
    assert any(isinstance(v, str) and "graph TD" in v for v in values)
    assert "PoC Flow" in values


def test_proposal_xlsx_project_name_override_reflected():
    content = generate_proposal_xlsx(_sample_ctx(project_name="Custom Export Name"))
    wb = load_workbook(BytesIO(content))
    summary = wb["Summary"]
    assert any(
        cell.value == "Custom Export Name" for row in summary.iter_rows() for cell in row
    )


def test_svg_to_png_bytes_empty_returns_none():
    assert svg_to_png_bytes("") is None
    assert svg_to_png_bytes("   ") is None


def test_svg_to_png_bytes_invalid_returns_none():
    assert svg_to_png_bytes("<not-valid-svg") is None


def test_docx_embeds_diagram_png_when_svg_present():
    ctx = _sample_ctx()
    # Minimal valid 1x1 PNG
    tiny_png = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc```\x00\x00"
        b"\x00\x03\x00\x01\x00\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    # Use a known-good tiny PNG from cairosvg of a simple SVG instead if needed
    from app.proposals.svg_raster import svg_to_png_bytes

    rendered = svg_to_png_bytes(
        '<svg xmlns="http://www.w3.org/2000/svg" width="80" height="40">'
        '<rect width="80" height="40" fill="#1E3A5F"/></svg>'
    )
    assert rendered
    import base64

    ctx["diagrams"] = [
        {
            "title": "Architecture Overview",
            "source": "graph TD\n  Client --> API",
            "png_base64": base64.b64encode(rendered).decode("ascii"),
        }
    ]
    content = generate_proposal_docx(ctx)
    document = Document(BytesIO(content))
    assert document.inline_shapes
    text = _docx_text(content)
    assert "Architecture Overview" in text
    assert "graph TD" not in text


def test_enrich_diagrams_copies_existing_svg(monkeypatch):
    from app.proposals.mermaid_render import enrich_diagrams_with_svg

    called = {"n": 0}

    def boom(_source: str):
        called["n"] += 1
        return b"\x89PNG\r\n\x1a\n"

    monkeypatch.setattr("app.proposals.mermaid_render.render_mermaid_png", boom)
    out = enrich_diagrams_with_svg(
        [{"title": "A", "source": "graph TD; A-->B", "png_base64": "already"}]
    )
    assert out[0]["png_base64"] == "already"
    assert called["n"] == 0


def test_proposal_xlsx_uses_presentation_theme_header_color():
    content = generate_proposal_xlsx(
        _sample_ctx(
            theme={
                "primary": "0F172A",
                "text_on_primary": "FFFFFF",
            },
            presentation={
                "theme_id": "modern-slate",
                "style_id": "spacious",
                "template_id": "executive-cover",
            },
        )
    )
    wb = load_workbook(BytesIO(content))
    summary = wb["Summary"]
    header = summary["A1"]
    assert header.fill.fgColor.rgb in {"000F172A", "0F172A", "FF0F172A"}
    values = [cell.value for row in summary.iter_rows() for cell in row]
    assert "modern-slate" in values
    assert "spacious" in values
    assert "executive-cover" in values


def test_proposal_docx_adds_cover_when_layout_requests_it():
    content = generate_proposal_docx(
        _sample_ctx(
            include_cover=True,
            layout={
                "layout": "executive_cover",
                "cover": True,
                "section_chrome": "minimal",
                "columns": 1,
            },
            theme={"primary": "0F172A", "text_on_primary": "FFFFFF"},
        )
    )
    text = _docx_text(content)
    assert "Portal Redesign" in text
    assert "Project Proposal Pack" in text


def test_proposal_markdown_includes_presentation_meta():
    content = generate_proposal_markdown(
        _sample_ctx(
            layout={"layout": "two_column", "cover": False, "section_chrome": "cards"},
            presentation={
                "theme_id": "warm-editorial",
                "style_id": "compact",
                "template_id": "two-column-summary",
            },
        )
    )
    text = content.decode("utf-8")
    assert "warm-editorial" in text
    assert "compact" in text
    assert "two-column-summary" in text


@pytest.mark.parametrize(
    ("size", "orientation", "expected"),
    [
        ("A4", "portrait", "size: A4 portrait;"),
        ("A4", "landscape", "size: A4 landscape;"),
        ("Letter", "portrait", "size: Letter portrait;"),
    ],
)
def test_proposal_html_uses_presentation_page_size(size, orientation, expected):
    ctx = _sample_ctx(
        include_cover=True,
        page={"size": size, "orientation": orientation, "css_size": f"{size} {orientation}"},
    )
    html = render_proposal_html(ctx)
    assert expected in html
    assert generate_proposal_pdf(ctx).startswith(b"%PDF")


@pytest.mark.parametrize(
    ("shape_type", "expected_tag"),
    [
        ("rectangle", "<rect"),
        ("line", "<line"),
        ("circle", "<circle"),
        ("ellipse", "<ellipse"),
        ("triangle", "<polygon"),
        ("polygon", "<polygon"),
    ],
)
def test_proposal_html_renders_backend_accent_svg_shape(shape_type, expected_tag):
    bundle = PresentationBundle(
        theme_id="theme",
        style_id="style",
        template_id="template",
        theme_tokens={"colors": {"primary": "17365D", "accent": "2563eb"}},
        layout={"cover": True},
        cover_design={
            "colors": {"background": "#f8fafc"},
            "accent_shapes": [_accent_shape(shape_type)],
        },
    )
    ctx = build_proposal_export_context(
        _proposal_for_export(),
        presentation=bundle,
        include_cover=True,
    )

    html = render_proposal_html(ctx)

    assert 'class="cover-accent-art"' in html
    assert expected_tag in html
    assert "#2563eb" in html
    assert "background: #f8fafc" in html
    assert "&lt;svg" not in html
    # WeasyPrint drops SVG fills when CSS forces width/height 100%.
    assert ".cover-accent-art svg { display: block; max-width: none; max-height: none; }" in html
    assert ".cover-accent-art svg { display: block; width: 100%; height: 100%; }" not in html


def test_proposal_context_renders_accents_background_and_warnings():
    bundle = PresentationBundle(
        theme_id="theme",
        style_id="style",
        template_id="template",
        theme_tokens={"colors": {"primary": "17365D", "accent": "2563EB"}},
        layout={"cover": True, "layout": "linear"},
        page={"size": "Legal", "orientation": "landscape"},
        cover_design={
            "colors": {"background": "#f8fafc"},
            "accent_shapes": [
                _accent_shape("rectangle", shape_id="back", z_index=1),
                _accent_shape("circle", shape_id="front", z_index=9),
                _accent_shape("ellipse", shape_id="hidden", visible=False),
            ],
        },
        accent_warnings=["Shape 3 was hidden during export"],
    )

    ctx = build_proposal_export_context(
        _proposal_for_export(),
        presentation=bundle,
        include_cover=True,
    )

    assert ctx["cover"]["background_color"] == "#f8fafc"
    assert 'width="355.6mm"' in ctx["cover"]["accent_svg"]
    assert 'height="215.9mm"' in ctx["cover"]["accent_svg"]
    assert ctx["cover"]["accent_svg"].index('data-accent-id="back"') < ctx["cover"][
        "accent_svg"
    ].index('data-accent-id="front"')
    assert 'data-accent-id="hidden"' not in ctx["cover"]["accent_svg"]
    assert "Shape 3 was hidden during export" in ctx["cover"]["warnings"]
    assert generate_proposal_pdf(ctx).startswith(b"%PDF")


def test_proposal_context_passes_resolved_chart_accent_to_gantt():
    proposal = _proposal_for_export()
    proposal.source_snapshot = {
        **proposal.source_snapshot,
        "gantt": {
            "project_start_date": "2026-08-01",
            "project_end_date": "2026-08-02",
            "tasks": [
                {
                    "name": "Design",
                    "phase": "design",
                    "start_date": "2026-08-01",
                    "end_date": "2026-08-02",
                }
            ],
        },
    }
    bundle = PresentationBundle(
        theme_id="theme",
        style_id="style",
        template_id="template",
        theme_tokens={
            "colors": {
                "accent": "C026D3",
                "chart": "0EA5E9",
            }
        },
    )

    ctx = build_proposal_export_context(proposal, presentation=bundle)

    assert 'fill="#0EA5E9"' in ctx["gantt_svg"]


def test_proposal_context_rejects_unsafe_cover_background():
    bundle = PresentationBundle(
        theme_id="theme",
        style_id="style",
        template_id="template",
        theme_tokens={"colors": {"primary": "17365D", "accent": "2563EB"}},
        layout={"cover": True},
        cover_design={"colors": {"background": "url(javascript:alert(1))"}},
    )

    ctx = build_proposal_export_context(
        _proposal_for_export(),
        presentation=bundle,
        include_cover=True,
    )

    assert ctx["cover"]["background_color"] == "#ffffff"


def test_proposal_html_renders_ordered_cover_fields_and_asset_regions():
    html = render_proposal_html(
        _sample_ctx(
            include_cover=True,
            cover={
                "fields": [
                    {"key": "title", "label": "Title", "value": "First", "emphasis": "title"},
                    {"key": "subtitle", "label": "Subtitle", "value": "Second"},
                ],
                "assets": [
                    {
                        "region": "background",
                        "url": "https://example.test/background.png",
                        "alt": "Cover background",
                    },
                    {
                        "region": "logo",
                        "url": "https://example.test/logo.png",
                        "alt": "Company logo",
                    },
                ],
                "design": {"alignment": "left", "padding_mm": 24},
                "warnings": [],
            },
        )
    )
    assert html.index("First") < html.index("Second")
    assert 'class="cover-asset cover-asset-background"' in html
    assert 'class="cover-asset cover-asset-logo"' in html


def test_proposal_html_renders_positioned_cover_field_geometry_and_typography():
    ctx = _sample_ctx(
        include_cover=True,
        cover={
            "fields": [
                {
                    "key": "title",
                    "label": "Title",
                    "value": "Positioned title",
                    "emphasis": "title",
                    "geometry": {
                        "x_pct": 12.5,
                        "y_pct": 8.25,
                        "width_pct": 62.0,
                        "height_pct": 14.5,
                        "z_index": 17,
                    },
                    "style": {
                        "font_family": 'A "Quoted"; Font',
                        "font_size_pt": 32.0,
                        "font_weight": 700,
                        "italic": True,
                        "color": "#123456",
                        "alignment": "center",
                        "line_height": 1.35,
                        "letter_spacing_pt": 1.5,
                        "opacity": 0.75,
                        "background_color": "#F0F2F4",
                        "padding_mm": 3.0,
                    },
                }
            ],
            "assets": [],
            "design": {},
            "warnings": [],
        },
    )

    html = render_proposal_html(ctx)

    assert "position:absolute" in html
    assert "left:12.5%" in html
    assert "top:8.25%" in html
    assert "width:62.0%" in html
    assert "height:14.5%" in html
    assert "z-index:17" in html
    assert "font-size:32.0pt" in html
    assert "font-weight:700" in html
    assert "font-style:italic" in html
    assert "color:#123456" in html
    assert "text-align:center" in html
    assert "line-height:1.35" in html
    assert "letter-spacing:1.5pt" in html
    assert "opacity:0.75" in html
    assert "background-color:#F0F2F4" in html
    assert "padding:3.0mm" in html
    assert "font-family:&#34;A \\&#34;Quoted\\&#34;; Font&#34;" in html
    assert "cover-field-styled" in html
    assert ".cover-field-styled .cover-field-value" in html
    # Positioned fields sit outside padded .cover-content (page-relative).
    content_start = html.index('class="cover-content"')
    field_start = html.index('data-cover-key="title"')
    assert field_start > content_start
    assert "position:absolute" in html[field_start : field_start + 200]
    assert generate_proposal_pdf(ctx).startswith(b"%PDF")


def test_proposal_html_keeps_legacy_fields_in_ordered_cover_flow():
    html = render_proposal_html(
        _sample_ctx(
            include_cover=True,
            cover={
                "fields": [
                    {"key": "title", "label": "Title", "value": "Legacy first"},
                    {"key": "client", "label": "Client", "value": "Legacy second"},
                ],
                "assets": [],
                "design": {},
                "warnings": [],
            },
        )
    )

    first_start = html.index('data-cover-key="title"')
    first_end = html.index("</div>", first_start)
    second_start = html.index('data-cover-key="client"')
    assert first_start < first_end < second_start
    assert "position:absolute" not in html[first_start:first_end]


def test_proposal_html_renders_cover_asset_geometry():
    html = render_proposal_html(
        _sample_ctx(
            include_cover=True,
            cover={
                "fields": [],
                "assets": [
                    {
                        "region": "logo",
                        "url": "https://example.test/logo.png",
                        "opacity": 0.6,
                        "geometry": {
                            "x_pct": 70.0,
                            "y_pct": 5.0,
                            "width_pct": 20.0,
                            "height_pct": 10.0,
                            "z_index": 23,
                            "rotation_deg": 15.0,
                        },
                    }
                ],
                "design": {},
                "warnings": [],
            },
        )
    )

    asset_start = html.index('class="cover-asset cover-asset-logo"')
    asset_end = html.index("/>", asset_start)
    asset_html = html[asset_start:asset_end]
    assert "opacity:0.6" in asset_html
    assert "left:70.0%" in asset_html
    assert "top:5.0%" in asset_html
    assert "width:20.0%" in asset_html
    assert "height:10.0%" in asset_html
    assert "z-index:23" in asset_html
    assert "right:auto" in asset_html
    assert "bottom:auto" in asset_html
    assert "max-width:none" in asset_html
    assert "max-height:none" in asset_html
    assert "transform:rotate(15.0deg)" in asset_html


@pytest.mark.parametrize(
    ("size", "orientation", "page_width_mm", "page_height_mm"),
    [
        ("A3", "portrait", 297, 420),
        ("A3", "landscape", 420, 297),
        ("A4", "portrait", 210, 297),
        ("A4", "landscape", 297, 210),
        ("Letter", "portrait", 215.9, 279.4),
        ("Letter", "landscape", 279.4, 215.9),
        ("Legal", "portrait", 215.9, 355.6),
        ("Legal", "landscape", 355.6, 215.9),
    ],
)
def test_proposal_pdf_cover_matches_printable_page_dimensions(
    size,
    orientation,
    page_width_mm,
    page_height_mm,
):
    html = render_proposal_html(
        _sample_ctx(
            include_cover=True,
            page={
                "size": size,
                "orientation": orientation,
                "css_size": f"{size} {orientation}",
            },
        )
    )

    assert f"--cover-page-width: {page_width_mm}mm" in html
    assert f"--cover-page-height: {page_height_mm}mm" in html
    assert "width: var(--cover-page-width);" in html
    assert "height: var(--cover-page-height);" in html
    assert "@page cover" in html
    assert "page: cover;" in html
    assert "min-height: 210mm" not in html
    assert "calc(var(--cover-page-width) - 16mm - 16mm)" not in html


def test_proposal_docx_gets_real_cover_page_and_fields():
    content = generate_proposal_docx(
        _sample_ctx(
            include_cover=True,
            cover={
                "fields": [
                    {"key": "title", "label": "Title", "value": "Executive Brief", "emphasis": "title"},
                    {"key": "client", "label": "Client", "value": "ACME Corp"},
                ],
                "assets": [],
                "design": {},
                "warnings": [],
            },
        )
    )
    document = Document(BytesIO(content))
    assert "Executive Brief" in _docx_text(content)
    assert "<w:br w:type=\"page\"" in document._element.xml


def test_proposal_docx_rasterizes_accent_before_cover_text(monkeypatch):
    accent_svg = (
        '<svg xmlns="http://www.w3.org/2000/svg" width="297mm" height="210mm">'
        '<rect width="297" height="210" fill="#2563eb" fill-opacity=".5"/></svg>'
    )
    png = svg_to_png_bytes(accent_svg)
    assert png is not None
    calls: list[tuple[str, float]] = []

    def rasterize(svg: str, *, scale: float = 1.5):
        calls.append((svg, scale))
        return png

    monkeypatch.setattr("app.exports.docx.svg_to_png_bytes", rasterize, raising=False)
    ctx = _sample_ctx(
        include_cover=True,
        gantt_svg="",
        page={"size": "A4", "orientation": "landscape", "css_size": "A4 landscape"},
        cover={
            "fields": [
                {
                    "key": "title",
                    "label": "Title",
                    "value": "Executive Brief",
                    "emphasis": "title",
                }
            ],
            "assets": [{"region": "logo", "url": "data:image/png;base64,unused"}],
            "design": {},
            "accent_svg": accent_svg,
            "warnings": [],
        },
    )

    content = generate_proposal_docx(ctx)

    document = Document(BytesIO(content))
    assert calls == [(accent_svg, 1.5)]
    assert len(document.inline_shapes) == 1
    assert document.inline_shapes[0].width > document.inline_shapes[0].height
    assert document._element.xml.index("<w:drawing>") < document._element.xml.index(
        "Executive Brief"
    )
    assert "Executive Brief" in _docx_text(content)
    assert ctx["cover"]["assets"][0]["region"] == "logo"


def test_proposal_docx_omits_failed_accent_and_adds_fidelity_warning(monkeypatch):
    monkeypatch.setattr(
        "app.exports.docx.svg_to_png_bytes",
        lambda _svg, **_kwargs: None,
        raising=False,
    )
    ctx = _sample_ctx(
        include_cover=True,
        gantt_svg="",
        cover={
            "fields": [
                {
                    "key": "title",
                    "label": "Title",
                    "value": "Fallback Brief",
                    "emphasis": "title",
                }
            ],
            "assets": [],
            "design": {},
            "accent_svg": "<svg xmlns='http://www.w3.org/2000/svg'/>",
            "warnings": [],
        },
    )

    content = generate_proposal_docx(ctx)

    document = Document(BytesIO(content))
    assert content[:2] == b"PK"
    assert len(document.inline_shapes) == 0
    assert "Fallback Brief" in _docx_text(content)
    assert (
        "DOCX uses flow-based Cover layout; exact positioning may differ from PDF."
        in ctx["cover"]["warnings"]
    )


def test_proposal_docx_documents_cover_geometry_as_flow_only():
    content = generate_proposal_docx(
        _sample_ctx(
            include_cover=True,
            cover={
                "fields": [
                    {
                        "key": "title",
                        "label": "Title",
                        "value": "Flow-only title",
                        "emphasis": "title",
                        "geometry": {
                            "x_pct": 70.0,
                            "y_pct": 80.0,
                            "width_pct": 20.0,
                            "z_index": 4,
                        },
                    }
                ],
                "assets": [],
                "design": {},
                "warnings": [],
            },
        )
    )

    document = Document(BytesIO(content))
    assert "Flow-only title" in _docx_text(content)
    assert "<wp:anchor" not in document._element.xml


def test_proposal_docx_applies_supported_cover_field_typography_in_flow():
    content = generate_proposal_docx(
        _sample_ctx(
            include_cover=True,
            cover={
                "fields": [
                    {
                        "key": "title",
                        "label": "Title",
                        "value": "Styled heading",
                        "emphasis": "title",
                        "geometry": {
                            "x_pct": 10.0,
                            "y_pct": 10.0,
                            "width_pct": 80.0,
                            "z_index": 4,
                        },
                        "style": {
                            "font_family": "Noto Sans JP",
                            "font_size_pt": 28.0,
                            "font_weight": 700,
                            "italic": True,
                            "color": "#123456",
                            "alignment": "center",
                            "line_height": 1.4,
                        },
                    }
                ],
                "assets": [],
                "design": {},
                "warnings": [],
            },
        )
    )

    document = Document(BytesIO(content))
    heading = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Styled heading")
    run = heading.runs[0]
    assert run.font.name == "Noto Sans JP"
    assert run.font.size.pt == pytest.approx(28.0)
    assert run.bold is True
    assert run.italic is True
    assert str(run.font.color.rgb) == "123456"
    assert heading.alignment == 1
    assert heading.paragraph_format.line_spacing == pytest.approx(1.4)


def test_markdown_and_xlsx_include_cover_metadata():
    ctx = _sample_ctx(
        include_cover=True,
        cover={
            "fields": [{"key": "subtitle", "label": "Subtitle", "value": "Transformation program"}],
            "assets": [{"region": "logo", "url": "https://example.test/logo.png"}],
            "design": {},
            "warnings": [],
        },
    )
    markdown = generate_proposal_markdown(ctx).decode("utf-8")
    assert "**Subtitle:** Transformation program" in markdown
    assert "Cover asset (logo)" in markdown

    workbook = load_workbook(BytesIO(generate_proposal_xlsx(ctx)))
    summary_values = [cell.value for row in workbook["Summary"].iter_rows() for cell in row]
    assert "Cover: Subtitle" in summary_values
    assert "Transformation program" in summary_values


@pytest.mark.asyncio
async def test_storage_backed_cover_asset_is_embedded_for_rendering():
    class Storage:
        async def read(self, path):
            assert path == "presentation-assets/template/background.png"
            return b"\x89PNG\r\n\x1a\n"

    ctx = {
        "cover": {
            "assets": [
                {
                    "region": "background",
                    "storage_path": "presentation-assets/template/background.png",
                    "url": "presentation-assets/template/background.png",
                }
            ]
        }
    }
    await _embed_cover_asset_data(ctx, Storage())
    assert ctx["cover"]["assets"][0]["url"].startswith("data:image/png;base64,")
