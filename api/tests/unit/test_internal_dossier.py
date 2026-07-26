from __future__ import annotations

import uuid
from datetime import datetime
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document
from openpyxl import load_workbook

from app.exports.internal_dossier import (
    build_internal_dossier_payload,
    build_internal_export_context,
    generate_internal_docx,
    generate_internal_markdown,
    generate_internal_pdf,
    generate_internal_xlsx,
    load_internal_export_parts,
)
from app.schemas.internal_dossier import InternalDossierResponse
from tests.unit.export_fixtures import sample_report_context


def test_export_context_includes_rate_card_and_proposal_markers():
    report = {
        "project_summary": {"project_name": "P"},
        "extracted": {"cost_drivers": [{"name": "x"}]},
    }
    rate_card = {
        "name": "RC",
        "version_number": 2,
        "settings": {"roles": [{"name": "Engineer", "hourly_rate": 10000}]},
    }
    proposals = [
        {
            "locale": "en",
            "status": "draft",
            "assessment": {"sections": []},
            "proposal_body": None,
            "poc": None,
        }
    ]

    ctx = build_internal_export_context(report, rate_card, proposals, locale="en")

    assert ctx["internal_banner"] == "INTERNAL — DO NOT DISTRIBUTE"
    assert ctx["rate_card"]["settings"]["roles"][0]["name"] == "Engineer"
    assert ctx["proposals"][0]["locale"] == "en"
    assert "cost_drivers" in ctx["report"]["extracted"]
    assert ctx["proposals_status"] == "present"


def test_export_context_marks_missing_proposal():
    ctx = build_internal_export_context({"project_summary": {}}, None, [], locale="en")

    assert ctx["proposals_status"] == "none"


@pytest.mark.asyncio
async def test_dossier_payload_loads_frozen_rate_card_and_proposals():
    estimate_id = uuid.uuid4()
    version_id = uuid.uuid4()
    rate_card_id = uuid.uuid4()
    proposal_id = uuid.uuid4()
    effective_date = datetime(2026, 7, 1, 12, 30)
    estimate = SimpleNamespace(
        id=estimate_id,
        project_name="Portal",
        client_name="ACME",
        status="calculated",
        calculation_result={"total_effort_hours": 10},
        rate_card_version_id=version_id,
        exports=[],
    )
    version = SimpleNamespace(
        id=version_id,
        rate_card_id=rate_card_id,
        version_number=3,
        created_at=effective_date,
        settings={"roles": [{"name": "Engineer", "hourly_rate": 10000}]},
    )
    rate_card = SimpleNamespace(id=rate_card_id, name="Standard")
    proposal = SimpleNamespace(
        id=proposal_id,
        locale="en",
        status="ready",
        include_poc=True,
        assessment={"sections": []},
        proposal_body={"sections": [{"title": "Approach"}]},
        poc={"sections": []},
    )
    rate_card_result = MagicMock()
    rate_card_result.one_or_none.return_value = (version, rate_card)
    proposals_result = MagicMock()
    proposals_result.scalars.return_value.all.return_value = [proposal]
    db = AsyncMock()
    db.execute = AsyncMock(side_effect=[rate_card_result, proposals_result])
    report = {"project_summary": {"project_name": "Portal"}}

    with (
        patch("app.exports.internal_dossier.build_report_context", return_value=report) as build,
        patch(
            "app.exports.internal_dossier.is_rate_card_stale_for_estimate",
            new=AsyncMock(return_value=True),
        ),
    ):
        payload = await build_internal_dossier_payload(db, estimate, locale="en")

    assert payload["estimate_id"] == str(estimate_id)
    assert payload["has_calculation"] is True
    assert payload["rate_card_stale"] is True
    assert payload["warnings"] == []
    assert payload["report"] == report
    assert payload["rate_card"] == {
        "rate_card_id": str(rate_card_id),
        "name": "Standard",
        "version_number": 3,
        "effective_date": effective_date.isoformat(),
        "settings": version.settings,
    }
    assert payload["proposals"] == [
        {
            "id": str(proposal_id),
            "locale": "en",
            "status": "ready",
            "include_poc": True,
            "assessment": proposal.assessment,
            "proposal_body": proposal.proposal_body,
            "poc": proposal.poc,
        }
    ]
    assert build.call_args.kwargs["rate_card_name"] == "Standard"
    assert build.call_args.kwargs["rate_card_version_number"] == 3
    assert build.call_args.kwargs["rate_card_effective_date"] == effective_date


@pytest.mark.asyncio
async def test_dossier_payload_warns_when_calculation_is_missing():
    estimate = SimpleNamespace(
        id=uuid.uuid4(),
        project_name="Draft",
        client_name="ACME",
        status="draft",
        calculation_result=None,
        rate_card_version_id=None,
        exports=[],
    )
    proposals_result = MagicMock()
    proposals_result.scalars.return_value.all.return_value = []
    db = AsyncMock()
    db.execute = AsyncMock(return_value=proposals_result)

    with (
        patch("app.exports.internal_dossier.build_report_context") as build,
        patch(
            "app.exports.internal_dossier.is_rate_card_stale_for_estimate",
            new=AsyncMock(return_value=False),
        ),
    ):
        payload = await build_internal_dossier_payload(db, estimate, locale="ja")

    build.assert_not_called()
    assert payload["report"] == {}
    assert payload["has_calculation"] is False
    assert payload["rate_card"] is None
    assert payload["warnings"]


@pytest.mark.asyncio
async def test_dossier_payload_treats_empty_calculation_as_present():
    estimate = SimpleNamespace(
        id=uuid.uuid4(),
        project_name="Calculated",
        client_name="ACME",
        status="calculated",
        calculation_result={},
        rate_card_version_id=None,
        exports=[],
    )
    proposals_result = MagicMock()
    proposals_result.scalars.return_value.all.return_value = []
    db = AsyncMock()
    db.execute = AsyncMock(return_value=proposals_result)

    with (
        patch("app.exports.internal_dossier.build_report_context", return_value={"labels": {}}) as build,
        patch(
            "app.exports.internal_dossier.is_rate_card_stale_for_estimate",
            new=AsyncMock(return_value=False),
        ),
    ):
        payload = await build_internal_dossier_payload(db, estimate, locale="en")

    build.assert_called_once()
    assert payload["has_calculation"] is True
    assert payload["warnings"] == []


@pytest.mark.asyncio
async def test_load_internal_export_parts_wraps_payload_for_generators():
    payload = {
        "locale": "en",
        "report": {"project_summary": {}},
        "rate_card": None,
        "proposals": [],
    }

    with patch(
        "app.exports.internal_dossier.build_internal_dossier_payload",
        new=AsyncMock(return_value=payload),
    ):
        context = await load_internal_export_parts(AsyncMock(), SimpleNamespace(), "en")

    assert context["internal_banner"] == "INTERNAL — DO NOT DISTRIBUTE"
    assert context["report"] == payload["report"]
    assert context["proposals_status"] == "none"


def test_internal_dossier_response_accepts_documented_shape():
    response = InternalDossierResponse(
        estimate_id=str(uuid.uuid4()),
        project_name="Portal",
        client_name="ACME",
        status="calculated",
        locale="en",
        has_calculation=True,
        rate_card_stale=False,
        warnings=[],
        report={"project_summary": {}},
        rate_card=None,
        proposals=[],
    )

    assert response.rate_card is None
    assert response.proposals == []


def test_internal_markdown_contains_banner_and_rate_card():
    ctx = build_internal_export_context(
        {"project_summary": {"project_name": "Alpha"}},
        {
            "name": "RC1",
            "settings": {
                "roles": [{"name": "PM", "hourly_rate": 1}],
                "setup_cost_items": [{"name": "Production setup", "amount": 250000}],
                "monthly_rc_items": [{"name": "Managed hosting", "amount": 50000}],
            },
        },
        [],
        locale="en",
    )
    ctx["internal_banner"] = ""
    md = generate_internal_markdown(ctx)
    assert "INTERNAL — DO NOT DISTRIBUTE" in md
    assert "Alpha" in md
    assert "PM" in md
    assert "Production setup" in md
    assert "Managed hosting" in md
    assert "none" in md.lower() or "No proposal" in md


def test_internal_pdf_is_pdf_and_html_has_banner():
    from app.exports.pdf import build_internal_dossier_html

    ctx = build_internal_export_context(
        {"project_summary": {"project_name": "Alpha"}, "labels": {}, "extracted": {}},
        {
            "name": "RC1",
            "settings": {
                "roles": [],
                "setup_cost_items": [{"name": "Production setup", "amount": 250000}],
                "monthly_rc_items": [{"name": "Managed hosting", "amount": 50000}],
            },
        },
        [],
        locale="en",
    )
    html = build_internal_dossier_html(ctx)
    assert "INTERNAL — DO NOT DISTRIBUTE" in html
    assert "Production setup" in html
    assert "Managed hosting" in html
    pdf = generate_internal_pdf(ctx)
    assert pdf.startswith(b"%PDF")


def _table_text(table) -> str:
    parts: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)
            for nested_table in cell.tables:
                parts.append(_table_text(nested_table))
    return "\n".join(parts)


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        parts.append(_table_text(table))
    return "\n".join(parts)


def _xlsx_text(content: bytes) -> str:
    wb = load_workbook(BytesIO(content))
    parts: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    parts.append(str(cell.value))
    return "\n".join(parts)


def _sample_internal_ctx(**overrides):
    defaults = dict(
        report=sample_report_context(),
        rate_card={
            "name": "RC1",
            "version_number": 2,
            "settings": {
                "roles": [{"name": "Senior Engineer", "hourly_rate": 12000}],
                "setup_cost_items": [{"name": "Production setup", "amount": 250000}],
                "monthly_rc_items": [{"name": "Managed hosting", "amount": 50000}],
            },
        },
        proposals=[
            {
                "locale": "en",
                "status": "ready",
                "assessment": {"sections": []},
                "proposal_body": {"sections": [{"title": "Approach"}]},
                "poc": None,
            }
        ],
        locale="en",
    )
    defaults.update(overrides)
    return build_internal_export_context(
        defaults["report"],
        defaults["rate_card"],
        defaults["proposals"],
        locale=defaults["locale"],
    )


def test_internal_docx_is_valid_zip_with_banner_and_rate_card():
    ctx = _sample_internal_ctx()

    content = generate_internal_docx(ctx)

    assert content[:2] == b"PK"
    text = _docx_text(content)
    assert "INTERNAL — DO NOT DISTRIBUTE" in text
    assert "Portal Redesign" in text
    assert "Senior Engineer" in text
    assert "Production setup" in text
    assert "Managed hosting" in text


def test_internal_docx_handles_missing_calculation_and_no_proposals():
    ctx = build_internal_export_context({"project_summary": {}}, None, [], locale="en")

    content = generate_internal_docx(ctx)

    assert content[:2] == b"PK"
    text = _docx_text(content)
    assert "INTERNAL — DO NOT DISTRIBUTE" in text


def test_internal_xlsx_is_valid_zip_with_banner_and_rate_card():
    ctx = _sample_internal_ctx()

    content = generate_internal_xlsx(ctx)

    assert content[:2] == b"PK"
    text = _xlsx_text(content)
    assert "INTERNAL — DO NOT DISTRIBUTE" in text
    assert "Portal Redesign" in text
    assert "Senior Engineer" in text
    assert "Production setup" in text
    assert "Managed hosting" in text


def test_internal_xlsx_handles_missing_calculation_and_no_proposals():
    ctx = build_internal_export_context({"project_summary": {}}, None, [], locale="en")

    content = generate_internal_xlsx(ctx)

    assert content[:2] == b"PK"
    text = _xlsx_text(content)
    assert "INTERNAL — DO NOT DISTRIBUTE" in text


def test_internal_markdown_includes_disclosure_appendix_fields():
    ctx = _sample_internal_ctx()

    md = generate_internal_markdown(ctx)

    assert "Internal Disclosure Appendix" in md
    assert "OAuth integration" in md  # cost_drivers
    assert "Third-party API changes" in md  # risks
    assert "High confidence on auth module." in md  # confidence_notes
    assert "Mobile support scope unclear" in md  # gaps


def test_internal_pdf_html_includes_disclosure_appendix_fields():
    from app.exports.pdf import build_internal_dossier_html

    ctx = _sample_internal_ctx()

    html = build_internal_dossier_html(ctx)

    assert "Internal Disclosure Appendix" in html
    assert "OAuth integration" in html  # cost_drivers
    assert "Third-party API changes" in html  # risks
    assert "High confidence on auth module." in html  # confidence_notes


def test_internal_pdf_html_includes_executive_cost_summary():
    from app.exports.pdf import build_internal_dossier_html

    ctx = _sample_internal_ctx()

    html = build_internal_dossier_html(ctx)

    assert "Executive Cost Summary" in html
    assert "Functional Requirements" in html
    assert "User authentication" in html


def test_internal_docx_includes_disclosure_appendix_fields():
    ctx = _sample_internal_ctx()

    content = generate_internal_docx(ctx)

    text = _docx_text(content)
    assert "Internal Disclosure Appendix" in text
    assert "OAuth integration" in text
    assert "Third-party API changes" in text
    assert "High confidence on auth module." in text


def test_internal_xlsx_includes_disclosure_appendix_fields():
    ctx = _sample_internal_ctx()

    content = generate_internal_xlsx(ctx)

    text = _xlsx_text(content)
    assert "OAuth integration" in text
    assert "Third-party API changes" in text
    assert "High confidence on auth module." in text


def test_internal_markdown_proposal_appendix_avoids_raw_dict_dump():
    ctx = build_internal_export_context(
        {"project_summary": {}},
        None,
        [
            {
                "locale": "en",
                "status": "ready",
                "assessment": {
                    "sections": [
                        {
                            "title": "Feasibility",
                            "body": "Looks feasible.",
                            "bullets": ["Clear scope"],
                            "rating": "green",
                        }
                    ]
                },
                "proposal_body": {"sections": [{"title": "Approach", "body": "Iterative delivery."}]},
                "poc": None,
            }
        ],
        locale="en",
    )

    md = generate_internal_markdown(ctx)

    assert "Feasibility" in md
    assert "Looks feasible." in md
    assert "Clear scope" in md
    assert "Rating: green" in md
    assert "{'sections'" not in md
    assert "{'title'" not in md


def test_internal_docx_proposal_appendix_avoids_raw_dict_dump():
    ctx = build_internal_export_context(
        {"project_summary": {}},
        None,
        [
            {
                "locale": "en",
                "status": "ready",
                "assessment": {
                    "sections": [
                        {"title": "Feasibility", "body": "Looks feasible.", "bullets": ["Clear scope"]}
                    ]
                },
                "proposal_body": None,
                "poc": None,
            }
        ],
        locale="en",
    )

    content = generate_internal_docx(ctx)

    text = _docx_text(content)
    assert "Feasibility" in text
    assert "Looks feasible." in text
    assert "Clear scope" in text
    assert "{'sections'" not in text


def test_internal_pdf_proposal_appendix_avoids_raw_dict_dump():
    from app.exports.pdf import build_internal_dossier_html

    ctx = build_internal_export_context(
        {"project_summary": {}, "labels": {}, "extracted": {}},
        None,
        [
            {
                "locale": "en",
                "status": "ready",
                "assessment": {
                    "sections": [
                        {"title": "Feasibility", "body": "Looks feasible.", "bullets": ["Clear scope"]}
                    ]
                },
                "proposal_body": None,
                "poc": None,
            }
        ],
        locale="en",
    )

    html = build_internal_dossier_html(ctx)

    assert "Feasibility" in html
    assert "Looks feasible." in html
    assert "Clear scope" in html
    assert "{&#39;sections&#39;" not in html
    assert "{'sections'" not in html
