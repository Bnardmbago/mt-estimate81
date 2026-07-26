from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.exports.markdown import (
    format_currency,
    format_effort_days,
    format_hours,
    format_person_days,
)

LOGO_PATH = Path(__file__).parent / "templates" / "assets" / "BI_logo.png"


def _document_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_heading(document: Document, text: str, *, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        table.rows[index].cells[0].text = label
        table.rows[index].cells[1].text = value
    document.add_paragraph()


def _add_data_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index + 1].cells
        for col_index, value in enumerate(row):
            cells[col_index].text = value
    document.add_paragraph()


def _add_bullet_list(document: Document, items: list[str]) -> None:
    if not items:
        document.add_paragraph("None")
        return
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def _add_subheading(document: Document, text: str) -> None:
    document.add_heading(text, level=3)


def _executive_pricing_rows(report_context: dict[str, Any]) -> list[tuple[str, str]]:
    labels = report_context["labels"]
    pricing = report_context.get("pricing_summary") or {}
    if pricing.get("has_discount"):
        return [
            (
                labels["development_cost_original"],
                format_currency(pricing["nrc_original_total_jpy"]),
            ),
            (labels["limited_time_discount"], pricing["discount_display"]),
            (
                labels["special_price"],
                f"{format_currency(pricing['nrc_discounted_total_jpy'])} {labels['excluding_tax']}",
            ),
        ]
    return [
        (
            labels["total_development_cost"],
            format_currency(report_context["executive_display"]["development_cost_jpy"]),
        ),
    ]


def build_report_document(report_context: dict[str, Any]) -> Document:
    """Build the standard report DOCX as a `Document` (not yet serialized to bytes).

    Exposed so callers (e.g. the internal dossier generator) can append
    additional sections to the same document instead of duplicating the
    report layout.
    """
    labels = report_context["labels"]
    project = report_context["project_summary"]
    executive_display = report_context["executive_display"]
    executive = report_context["executive_summary"]
    effort = report_context["effort_summary"]
    calculation = report_context["calculation"]
    extracted = report_context["extracted"]
    gantt = report_context.get("gantt") or {}

    document = Document()
    _add_heading(document, labels["title"])

    _add_heading(document, labels["project_summary"], level=2)
    _add_key_value_table(
        document,
        [
            (labels["project_name"], project["project_name"]),
            (labels["estimate_type"], project["estimate_type"] or labels["none"]),
            (labels["client_name"], project["client_name"]),
            (labels["estimate_id"], project["estimate_id"]),
            (labels["export_revision"], str(project["export_revision"])),
            (labels["generated_date"], project["generated_date"]),
            (labels["estimate_creator"], project["estimate_creator"]),
        ],
    )

    _add_heading(document, labels["executive_cost_summary"], level=2)
    _add_key_value_table(
        document,
        [
            *_executive_pricing_rows(report_context),
            (
                labels["maintenance_cost_monthly_annual"],
                executive_display["maintenance_cost_display"],
            ),
            (
                labels["development_period"],
                executive_display["development_period_display"],
            ),
        ],
    )
    pricing = report_context.get("pricing_summary") or {}
    if pricing.get("has_discount") and pricing.get("campaign_terms"):
        document.add_paragraph(f"*{pricing.get('campaign_terms_title', labels['campaign_terms_title'])}")
        document.add_paragraph(pricing["campaign_terms"])

    _add_heading(document, labels["questionnaire"], level=2)
    questionnaire_sections = report_context.get("questionnaire_sections") or []
    if questionnaire_sections:
        for section in questionnaire_sections:
            _add_subheading(document, section["title"])
            _add_key_value_table(
                document,
                [(field["label"], field["value"]) for field in section["fields"]],
            )
    else:
        document.add_paragraph(labels["none"])

    requirement_sections = [
        (labels["functional_requirements"], extracted.get("functional_requirements") or []),
        (labels["non_functional_requirements"], extracted.get("non_functional_requirements") or []),
        (labels["modules"], extracted.get("modules") or []),
        (labels["user_roles"], extracted.get("user_roles") or []),
        (labels["external_systems"], extracted.get("external_systems") or []),
    ]
    has_requirements = any(items for _, items in requirement_sections)
    if has_requirements:
        for title, items in requirement_sections:
            if items:
                _add_heading(document, title, level=2)
                _add_bullet_list(document, items)
    else:
        _add_heading(document, labels["functional_requirements"], level=2)
        document.add_paragraph(labels["none"])

    _add_heading(document, labels["feature_items"], level=2)
    _add_data_table(
        document,
        [
            labels["feature_name"],
            labels["feature_description"],
            labels["feature_phase"],
            labels["feature_role"],
            labels["feature_hours"],
            labels["feature_days"],
        ],
        [
            [
                item["name"],
                item["description"],
                item["phase"],
                item["role"],
                format_hours(item["hours"]),
                format_effort_days(item["hours"]),
            ]
            for item in report_context["feature_items"]
        ],
    )

    _add_heading(document, labels["effort_summary"], level=2)
    _add_key_value_table(
        document,
        [
            (labels["total_hours"], format_hours(effort["total_hours"])),
            (labels["total_days"], format_effort_days(effort["total_hours"])),
            (
                labels["estimated_duration"],
                f"{format_person_days(effort['estimated_duration_days'])} {labels['days']}",
            ),
        ],
    )

    if gantt.get("tasks"):
        _add_heading(document, labels["gantt_title"], level=2)
        document.add_paragraph(labels["gantt_assumption"])
        _add_key_value_table(
            document,
            [
                (labels["gantt_project_start"], gantt["project_start_date"]),
                (labels["gantt_project_end"], gantt["project_end_date"]),
                (labels["gantt_total_working_days"], str(gantt["total_working_days"])),
            ],
        )

    _add_heading(document, labels["nrc_detailed"], level=2)
    nrc_rows = [
        [row["category"], row["item"], format_currency(row["cost_jpy"])]
        for row in calculation["nrc_line_items"]
    ]
    nrc_rows.append(
        [labels["nrc_total"], "", format_currency(executive["nrc_total_jpy"])]
    )
    _add_data_table(document, [labels["category"], labels["item"], labels["cost"]], nrc_rows)

    _add_heading(document, labels["rc_detailed"], level=2)
    rc_breakdown = report_context.get("rc_breakdown") or {}
    rc_rows = []
    for row in rc_breakdown.get("line_items") or []:
        item_label = row.get("service_description") or (
            labels["maintenance"] if row.get("is_maintenance") else row["item"]
        )
        rc_rows.append(
            [
                row["category"],
                item_label,
                format_currency(row["monthly_jpy"]),
                format_currency(row["annual_jpy"]),
            ]
        )
    rc_rows.append(
        [
            labels["monthly_total"],
            "",
            format_currency(rc_breakdown.get("monthly_total_jpy", 0)),
            "",
        ]
    )
    rc_rows.append(
        [
            labels["annual_total"],
            "",
            "",
            format_currency(rc_breakdown.get("annual_total_jpy", 0)),
        ]
    )
    _add_data_table(
        document,
        [labels["category"], labels["service_description"], labels["monthly"], labels["annual"]],
        rc_rows,
    )

    _add_heading(document, labels["estimate_exclusions"], level=2)
    _add_bullet_list(document, extracted.get("estimate_exclusions") or [labels["none"]])

    _add_heading(document, labels["approval"], level=2)
    document.add_paragraph(f"{labels['prepared_by']}: {labels['prepared_by_value']}")
    document.add_paragraph(f"{labels['reviewed_by']}:")
    document.add_paragraph(f"{labels['approved_by']}:")
    document.add_paragraph(f"{labels['approval_date']}:")

    return document


def generate_report_docx(report_context: dict[str, Any]) -> bytes:
    return _document_bytes(build_report_document(report_context))


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    font_size: Pt | None = None,
) -> None:
    cell.text = text
    if not cell.paragraphs[0].runs:
        return
    run = cell.paragraphs[0].runs[0]
    if bold:
        run.bold = True
    if font_size is not None:
        run.font.size = font_size


def generate_quotation_docx(quotation_context: dict[str, Any]) -> bytes:
    """Backward-compatible alias for unified formal quotation DOCX generation."""
    return generate_quotation_formal_docx(quotation_context)


def generate_quotation_formal_docx(quotation_context: dict[str, Any]) -> bytes:
    labels = quotation_context["labels"]
    company = quotation_context["company"]
    locale = quotation_context["locale"]
    colon = "：" if locale == "ja" else ": "

    document = Document()

    header = document.add_table(rows=3, cols=2)
    header.autofit = True

    title_cell = header.rows[0].cells[0]
    right_cell = header.rows[0].cells[1]
    title = title_cell.paragraphs[0].add_run(labels["title"])
    title.bold = True
    title.font.size = Pt(18)

    date_paragraph = right_cell.paragraphs[0]
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_paragraph.add_run(quotation_context["issue_date"])
    date_run.bold = True
    date_run.font.size = Pt(12)
    right_cell.add_paragraph()
    right_cell.add_paragraph(
        f"{labels['quote_number']}{colon}{quotation_context['quote_number']}"
    )
    right_cell.add_paragraph(
        f"{labels['registration_number']}{colon}{quotation_context['registration_number']}"
    )
    logo_bytes = quotation_context.get("logo_bytes")
    logo_ext = (quotation_context.get("logo_ext") or "").lower()
    if logo_bytes and logo_ext in {"png", "jpg", "jpeg", "webp"}:
        logo_paragraph = right_cell.add_paragraph()
        logo_paragraph.add_run().add_picture(BytesIO(logo_bytes), width=Inches(1.6))
    elif LOGO_PATH.exists():
        logo_paragraph = right_cell.add_paragraph()
        logo_paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.6))
    contact_person = company.get("contact_person") or ""
    right_cell.add_paragraph(f"{labels['contact_person']}{colon}{contact_person}")
    right_cell.add_paragraph()
    postal_prefix = "〒" if locale == "ja" else "Postal code "
    right_cell.add_paragraph(f"{postal_prefix}{company['postal_code']}")
    for line in company["address_lines"]:
        right_cell.add_paragraph(line)
    right_cell.add_paragraph()
    right_cell.add_paragraph(f"{labels['tel']}{colon}{company['tel']}")
    mail_sep = "：" if locale == "ja" else ": "
    right_cell.add_paragraph(f"{labels['mail']}{mail_sep}{company['email']}")
    right_cell.add_paragraph()

    right_cell.merge(header.rows[1].cells[1])
    right_cell.merge(header.rows[2].cells[1])

    client_cell = header.rows[1].cells[0]
    client_name_paragraph = client_cell.paragraphs[0]
    client_name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_name_paragraph.paragraph_format.space_before = Pt(28)
    client_name_run = client_name_paragraph.add_run(quotation_context["client_name"])
    client_name_run.bold = True
    client_name_run.font.size = Pt(14)

    intro_cell = header.rows[2].cells[0]
    intro_cell.paragraphs[0].add_run(quotation_context["intro"])

    document.add_paragraph()

    total_table = document.add_table(rows=1, cols=1)
    total_table.style = "Table Grid"
    total_cell = total_table.rows[0].cells[0]
    total_paragraph = total_cell.paragraphs[0]
    total_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    amount_run = total_paragraph.add_run(
        format_currency(quotation_context["grand_total_jpy"])
    )
    amount_run.bold = True
    amount_run.font.size = Pt(14)

    document.add_paragraph()

    line_items = quotation_context["line_items"]
    items_table = document.add_table(rows=1 + len(line_items), cols=4)
    items_table.style = "Table Grid"
    headers = [labels["item"], labels["unit"], labels["unit_price"], labels["subtotal_col"]]
    for index, header in enumerate(headers):
        _set_cell_text(items_table.rows[0].cells[index], header, bold=True)

    for row_index, row in enumerate(line_items):
        cells = items_table.rows[row_index + 1].cells
        name_text = row["name"]
        if row.get("description"):
            name_text = f"{name_text}\n{row['description']}"
        cells[0].text = name_text
        cells[1].text = str(row["quantity"]) if row.get("display_quantity") else ""
        if row.get("kind") == "discount":
            cells[2].text = ""
            cells[3].text = row.get("discount_display", "")
        else:
            cells[2].text = f"({format_currency(row['unit_price_jpy'])})"
            cells[3].text = f"({format_currency(row['subtotal_jpy'])})"

    document.add_paragraph()

    totals_layout = document.add_table(rows=1, cols=2)
    totals_layout.autofit = True

    notes_cell = totals_layout.rows[0].cells[0]
    notes_heading = notes_cell.paragraphs[0].add_run(labels["notes_heading"])
    notes_heading.bold = True
    for item in quotation_context.get("remarks_items") or []:
        notes_cell.add_paragraph(f"・{item}")

    totals = totals_layout.rows[0].cells[1].add_table(rows=3, cols=2)
    totals.style = "Table Grid"
    totals_rows = [
        (labels["subtotal"], f"{quotation_context['subtotal_jpy']:,}"),
        (
            quotation_context["tax_with_rate_label"],
            f"{quotation_context['tax_jpy']:,}",
        ),
        (
            labels["grand_total"],
            f"{quotation_context['grand_total_jpy']:,}",
        ),
    ]
    for index, (label, value) in enumerate(totals_rows):
        _set_cell_text(totals.rows[index].cells[0], label, bold=True, font_size=Pt(11))
        _set_cell_text(totals.rows[index].cells[1], value, font_size=Pt(11))

    document.add_paragraph()
    document.add_paragraph(labels["bank_details"]).runs[0].bold = True
    bank_paragraph = document.add_paragraph()
    bank_run = bank_paragraph.add_run(quotation_context["bank_details"])
    bank_run.font.size = Pt(11)

    return _document_bytes(document)
