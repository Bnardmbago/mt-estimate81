from __future__ import annotations

import uuid
from datetime import UTC, datetime
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.estimates import service as estimate_service
from app.estimates.rate_card_stale import is_rate_card_stale_for_estimate
from app.exports.export_i18n import localize_role
from app.exports.markdown import generate_markdown
from app.exports.report_context import _enrich_role_breakdown, build_report_context
from app.models.estimate import Estimate
from app.models.proposal import Proposal
from app.models.rate_card import RateCard, RateCardVersion
from app.models.user import User
from app.presentation.resolver import PresentationBundle
from app.proposals.export_pack_content import brief_field_rows, section_rows


INTERNAL_BANNER = "INTERNAL — DO NOT DISTRIBUTE"
MISSING_CALCULATION_WARNING = "Calculation result is unavailable."
MISSING_RATE_CARD_WARNING = "Rate card version is unavailable."


def _restore_full_role_breakdown_for_dossier(
    report: dict[str, Any],
    estimate: Estimate,
    *,
    locale: str,
) -> None:
    """Keep every frozen rate-card role (incl. 0h) so Estimate tab matches Rate Card tab.

    ``build_report_context`` filters inactive roles for client exports; the dossier UI
    should list the same roles as the attached rate card version.
    """
    calculation = report.get("calculation")
    raw = (estimate.calculation_result or {}).get("role_breakdown") or []
    if not calculation or not raw:
        return

    total_days = float((estimate.calculation_result or {}).get("total_effort_days") or 0)
    estimated_duration_days = float(
        (estimate.calculation_result or {}).get("estimated_duration_days") or total_days or 0
    )
    enriched = _enrich_role_breakdown(
        list(raw),
        estimated_duration_days=estimated_duration_days,
        total_days=total_days,
    )
    if locale == "ja":
        enriched = [
            {**row, "role": localize_role(str(row.get("role") or ""), locale)}
            for row in enriched
        ]
    calculation["role_breakdown"] = enriched


def build_internal_export_context(
    report: dict[str, Any],
    rate_card: dict[str, Any] | None,
    proposals: list[dict[str, Any]],
    *,
    locale: str,
) -> dict[str, Any]:
    return {
        "locale": locale,
        "internal_banner": INTERNAL_BANNER,
        "report": report,
        "rate_card": rate_card,
        "proposals": proposals,
        "proposals_status": "none" if not proposals else "present",
    }


async def _load_rate_card(
    db: AsyncSession,
    rate_card_version_id: Any,
) -> tuple[dict[str, Any] | None, datetime | None]:
    if not rate_card_version_id:
        return None, None

    result = await db.execute(
        select(RateCardVersion, RateCard)
        .join(RateCard, RateCard.id == RateCardVersion.rate_card_id)
        .where(RateCardVersion.id == rate_card_version_id)
    )
    row = result.one_or_none()
    if row is None:
        return None, None

    version, rate_card = row
    effective_date = version.created_at
    return (
        {
            "rate_card_id": str(version.rate_card_id) if version.rate_card_id else None,
            "name": rate_card.name,
            "version_number": version.version_number,
            "effective_date": effective_date.isoformat() if effective_date else None,
            "settings": version.settings or {},
        },
        effective_date,
    )


async def _load_proposals(
    db: AsyncSession,
    estimate_id: Any,
) -> list[dict[str, Any]]:
    result = await db.execute(
        select(Proposal)
        .where(Proposal.estimate_id == estimate_id)
        .order_by(Proposal.created_at.asc())
    )
    proposals = result.scalars().all()
    return [
        {
            "id": str(proposal.id),
            "locale": proposal.locale,
            "status": proposal.status,
            "include_poc": proposal.include_poc,
            "assessment": proposal.assessment,
            "proposal_body": proposal.proposal_body,
            "poc": proposal.poc,
        }
        for proposal in proposals
    ]


async def build_internal_dossier_payload(
    db: AsyncSession,
    estimate: Estimate,
    *,
    locale: str,
    presentation: PresentationBundle | None = None,
    include_cover: bool | None = None,
    cover_values: dict[str, Any] | None = None,
) -> dict[str, Any]:
    warnings: list[str] = []
    rate_card, rate_card_effective_date = await _load_rate_card(
        db,
        estimate.rate_card_version_id,
    )
    if estimate.rate_card_version_id and rate_card is None:
        warnings.append(MISSING_RATE_CARD_WARNING)

    has_calculation = estimate.calculation_result is not None
    if has_calculation:
        report = build_report_context(
            estimate,
            locale,
            generated_at=datetime.now(UTC),
            rate_card_name=rate_card["name"] if rate_card else None,
            rate_card_version_number=rate_card["version_number"] if rate_card else None,
            rate_card_effective_date=rate_card_effective_date,
            export_revision=1,
            presentation=presentation,
            include_cover=include_cover,
            cover_values=cover_values,
        )
        _restore_full_role_breakdown_for_dossier(report, estimate, locale=locale)
    else:
        report = {}
        warnings.append(MISSING_CALCULATION_WARNING)

    proposals = await _load_proposals(db, estimate.id)
    return {
        "estimate_id": str(estimate.id),
        "project_name": estimate.project_name,
        "client_name": estimate.client_name,
        "status": estimate.status,
        "locale": locale,
        "has_calculation": has_calculation,
        "rate_card_stale": await is_rate_card_stale_for_estimate(db, estimate),
        "warnings": warnings,
        "report": report,
        "rate_card": rate_card,
        "proposals": proposals,
    }


async def get_internal_dossier(
    db: AsyncSession,
    estimate_id: uuid.UUID,
    admin: User,
) -> dict[str, Any]:
    estimate = await estimate_service.get_estimate_for_user(db, estimate_id, admin)
    return await build_internal_dossier_payload(db, estimate, locale=estimate.locale)


async def load_internal_export_parts(
    db: AsyncSession,
    estimate: Estimate,
    locale: str,
    *,
    presentation: PresentationBundle | None = None,
    include_cover: bool | None = None,
    cover_values: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Return build_internal_export_context(...) ready for generators."""
    payload = await build_internal_dossier_payload(
        db,
        estimate,
        locale=locale,
        presentation=presentation,
        include_cover=include_cover,
        cover_values=cover_values,
    )
    return build_internal_export_context(
        payload["report"],
        payload["rate_card"],
        payload["proposals"],
        locale=locale,
    )


def _fallback_report_markdown(report: dict[str, Any]) -> str:
    project_summary = report.get("project_summary") or {}
    lines = ["## Project Summary", ""]
    for label, key in (
        ("Project Name", "project_name"),
        ("Client", "client_name"),
        ("Estimate ID", "estimate_id"),
    ):
        value = project_summary.get(key)
        if value:
            lines.append(f"- {label}: {value}")
    if len(lines) == 2:
        lines.append(f"- {MISSING_CALCULATION_WARNING}")
    return "\n".join(lines)


def _rate_card_markdown(rate_card: dict[str, Any] | None) -> str:
    lines = ["## Rate Card Appendix", ""]
    if not rate_card:
        lines.append("- none")
        return "\n".join(lines)

    lines.append(f"- Rate Card Name: {rate_card.get('name', '')}")
    version_number = rate_card.get("version_number")
    if version_number is not None:
        lines.append(f"- Version: {version_number}")

    settings = rate_card.get("settings") or {}

    roles = settings.get("roles") or []
    if roles:
        lines.append("")
        lines.append("| Role | Hourly Rate |")
        lines.append("|---|---:|")
        for role in roles:
            lines.append(f"| {role.get('name', '')} | {role.get('hourly_rate', '')} |")

    setup_cost_items = settings.get("setup_cost_items") or settings.get("nrc_items") or []
    if setup_cost_items:
        lines.append("")
        lines.append("| NRC Item | Cost |")
        lines.append("|---|---:|")
        for item in setup_cost_items:
            name = item.get("name") or item.get("item", "")
            cost = item.get("amount", item.get("cost_jpy", item.get("cost", "")))
            lines.append(f"| {name} | {cost} |")

    monthly_rc_items = settings.get("monthly_rc_items") or settings.get("rc_items") or []
    if monthly_rc_items:
        lines.append("")
        lines.append("| RC Item | Monthly |")
        lines.append("|---|---:|")
        for item in monthly_rc_items:
            name = item.get("name") or item.get("item", "")
            monthly = item.get("amount", item.get("monthly_jpy", item.get("monthly", "")))
            lines.append(f"| {name} | {monthly} |")

    return "\n".join(lines)


# Disclosure fields pulled from report["extracted"] that client-facing reports omit.
# Each entry renders as a bulleted list under the Internal Disclosure Appendix.
_DISCLOSURE_LIST_FIELDS: list[tuple[str, str]] = [
    ("risks", "Risks"),
    ("gaps", "Gaps"),
    ("confidence_factors", "Confidence Factors"),
    ("missing_inputs", "Missing Inputs"),
    ("recommendations", "Recommendations"),
    ("estimation_warnings", "Estimation Warnings"),
    ("assumption_risks", "Assumption Risks"),
    ("estimate_exclusions", "Estimate Exclusions"),
]

# Labels for the project brief fields shared across assessment/proposal/poc blobs.
_BRIEF_LABELS: dict[str, str] = {
    "brief_project_name": "Project Name",
    "brief_description": "Description",
    "brief_business_problem": "Business Problem",
    "brief_target_users": "Target Users",
    "brief_technology_stack": "Technology Stack",
    "brief_constraints": "Constraints",
}


def _cost_driver_rows(cost_drivers: list[Any] | None) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    for driver in cost_drivers or []:
        if isinstance(driver, dict):
            name = str(driver.get("name") or driver.get("driver") or "")
            impact = driver.get("impact_jpy", driver.get("impact", ""))
            rows.append((name, str(impact) if impact != "" else ""))
        else:
            rows.append((str(driver), ""))
    return rows


def _disclosure_markdown(report: dict[str, Any]) -> str:
    """Render the Internal Disclosure Appendix: fields client reports omit."""
    extracted = report.get("extracted") or {}
    lines = ["## Internal Disclosure Appendix", ""]

    lines.append("### Cost Drivers")
    lines.append("")
    cost_driver_rows = _cost_driver_rows(extracted.get("cost_drivers"))
    if cost_driver_rows:
        lines.append("| Driver | Impact (JPY) |")
        lines.append("|---|---:|")
        for name, impact in cost_driver_rows:
            lines.append(f"| {name} | {impact} |")
    else:
        lines.append("- none")
    lines.append("")

    for key, label in _DISCLOSURE_LIST_FIELDS:
        items = extracted.get(key) or []
        lines.append(f"### {label}")
        lines.append("")
        if items:
            lines.extend(f"- {item}" for item in items)
        else:
            lines.append("- none")
        lines.append("")

    lines.append("### Confidence")
    lines.append("")
    score = extracted.get("confidence_score")
    lines.append(f"- Score: {score if score is not None else 'none'}")
    lines.append(f"- Notes: {extracted.get('confidence_notes') or 'none'}")
    lines.append("")

    lines.append("### Questionnaire Appendix")
    lines.append("")
    questionnaire_sections = report.get("questionnaire_sections") or []
    if questionnaire_sections:
        for section in questionnaire_sections:
            lines.append(f"#### {section.get('title', '')}")
            lines.append("")
            for field in section.get("fields") or []:
                lines.append(f"- **{field.get('label')}:** {field.get('value')}")
            lines.append("")
    else:
        lines.append("- none")
        lines.append("")

    return "\n".join(lines)


def _proposal_part_markdown(title: str, blob: dict[str, Any] | None) -> list[str]:
    """Render an assessment/proposal_body/poc blob as structured sections (not str(dict))."""
    lines = [f"#### {title}"]
    lines.append("")
    if not blob:
        lines.append("- none")
        lines.append("")
        return lines

    brief_rows = brief_field_rows(blob.get("project_brief"), _BRIEF_LABELS)
    if brief_rows:
        lines.append("##### Project Brief")
        lines.append("")
        for label, value in brief_rows:
            lines.append(f"- **{label}:** {value}")
        lines.append("")

    for row in section_rows(blob):
        if row["title"]:
            lines.append(f"##### {row['title']}")
            lines.append("")
        if row["body"]:
            lines.append(str(row["body"]))
            lines.append("")
        for bullet in row["bullets"]:
            lines.append(f"- {bullet}")
        if row["bullets"]:
            lines.append("")
        if row.get("rating"):
            lines.append(f"_Rating: {row['rating']}_")
            lines.append("")

    if not brief_rows and not blob.get("sections"):
        lines.append("- none")
        lines.append("")

    return lines


def _proposals_markdown(proposals: list[dict[str, Any]]) -> str:
    lines = ["## Proposal Appendix", ""]
    if not proposals:
        lines.append("No proposal (none)")
        return "\n".join(lines)

    for proposal in proposals:
        lines.append(f"### Proposal ({proposal.get('locale', '')}) — {proposal.get('status', '')}")
        lines.append("")
        lines.extend(_proposal_part_markdown("Assessment", proposal.get("assessment")))
        lines.extend(_proposal_part_markdown("Proposal Body", proposal.get("proposal_body")))
        lines.extend(_proposal_part_markdown("POC", proposal.get("poc")))

    return "\n".join(lines)


def generate_internal_markdown(ctx: dict[str, Any]) -> str:
    """Build the internal Markdown dossier: banner + report + disclosure + rate card + proposals."""
    report = ctx.get("report") or {}

    sections = [f"# {INTERNAL_BANNER}", ""]

    if report.get("labels"):
        sections.append(generate_markdown(report))
    else:
        sections.append(_fallback_report_markdown(report))

    sections.append("")
    sections.append(_disclosure_markdown(report))
    sections.append("")
    sections.append(_rate_card_markdown(ctx.get("rate_card")))
    sections.append("")
    sections.append(_proposals_markdown(ctx.get("proposals") or []))

    return "\n".join(sections)


def generate_internal_pdf(ctx: dict[str, Any]) -> bytes:
    """Render the internal PDF dossier (banner + report + rate card + proposals)."""
    from app.exports.pdf import generate_internal_dossier_pdf

    return generate_internal_dossier_pdf(ctx)


def _rate_card_role_rows(rate_card: dict[str, Any] | None) -> list[tuple[str, str]]:
    if not rate_card:
        return [("Rate Card", "none")]
    rows = [("Rate Card Name", str(rate_card.get("name", "")))]
    version_number = rate_card.get("version_number")
    if version_number is not None:
        rows.append(("Version", str(version_number)))
    return rows


def _rate_card_role_table(rate_card: dict[str, Any] | None) -> tuple[list[str], list[list[str]]] | None:
    roles = (rate_card or {}).get("settings", {}).get("roles") or []
    if not roles:
        return None
    return (
        ["Role", "Hourly Rate"],
        [[str(role.get("name", "")), str(role.get("hourly_rate", ""))] for role in roles],
    )


def _rate_card_item_table(
    settings: dict[str, Any],
    *,
    keys: tuple[str, ...],
    headers: list[str],
) -> tuple[list[str], list[list[str]]] | None:
    items: list[dict[str, Any]] = []
    for key in keys:
        items = settings.get(key) or []
        if items:
            break
    if not items:
        return None
    rows = []
    for item in items:
        name = item.get("name") or item.get("item", "")
        value = item.get("amount", item.get("cost_jpy", item.get("monthly_jpy", item.get("cost", item.get("monthly", "")))))
        rows.append([str(name), str(value)])
    return (headers, rows)


def _add_disclosure_docx(document: Any, report: dict[str, Any]) -> None:
    """DOCX rendering of the Internal Disclosure Appendix: fields client reports omit."""
    from app.exports.docx import _add_bullet_list, _add_data_table, _add_key_value_table, _add_subheading

    extracted = report.get("extracted") or {}

    _add_subheading(document, "Cost Drivers")
    cost_driver_rows = _cost_driver_rows(extracted.get("cost_drivers"))
    if cost_driver_rows:
        _add_data_table(document, ["Driver", "Impact (JPY)"], [list(row) for row in cost_driver_rows])
    else:
        document.add_paragraph("None")

    for key, label in _DISCLOSURE_LIST_FIELDS:
        _add_subheading(document, label)
        _add_bullet_list(document, [str(item) for item in extracted.get(key) or []])

    _add_subheading(document, "Confidence")
    score = extracted.get("confidence_score")
    _add_key_value_table(
        document,
        [
            ("Score", str(score) if score is not None else "None"),
            ("Notes", str(extracted.get("confidence_notes") or "None")),
        ],
    )

    _add_subheading(document, "Questionnaire Appendix")
    questionnaire_sections = report.get("questionnaire_sections") or []
    if questionnaire_sections:
        for section in questionnaire_sections:
            document.add_heading(str(section.get("title") or ""), level=4)
            _add_key_value_table(
                document,
                [(field.get("label", ""), str(field.get("value", ""))) for field in section.get("fields") or []],
            )
    else:
        document.add_paragraph("None")


def _add_proposal_part_docx(document: Any, title: str, blob: dict[str, Any] | None) -> None:
    """Render an assessment/proposal_body/poc blob as structured sections (not str(dict))."""
    from app.exports.docx import _add_bullet_list, _add_key_value_table, _add_subheading

    _add_subheading(document, title)
    if not blob:
        document.add_paragraph("None")
        return

    brief_rows = brief_field_rows(blob.get("project_brief"), _BRIEF_LABELS)
    if brief_rows:
        document.add_heading("Project Brief", level=4)
        _add_key_value_table(document, brief_rows)

    rows = section_rows(blob)
    for row in rows:
        if row["title"]:
            document.add_heading(row["title"], level=4)
        if row["body"]:
            document.add_paragraph(str(row["body"]))
        if row["bullets"]:
            _add_bullet_list(document, row["bullets"])
        if row.get("rating"):
            document.add_paragraph(f"Rating: {row['rating']}")

    if not brief_rows and not rows:
        document.add_paragraph("None")


def _add_fallback_report_docx(document: Any, report: dict[str, Any]) -> None:
    project_summary = report.get("project_summary") or {}
    document.add_heading("Project Summary", level=2)
    from app.exports.docx import _add_key_value_table

    rows = [
        (label, str(project_summary[key]))
        for label, key in (
            ("Project Name", "project_name"),
            ("Client", "client_name"),
            ("Estimate ID", "estimate_id"),
        )
        if project_summary.get(key)
    ]
    if not rows:
        document.add_paragraph(MISSING_CALCULATION_WARNING)
        return
    _add_key_value_table(document, rows)


def generate_internal_docx(ctx: dict[str, Any]) -> bytes:
    """Render the internal DOCX dossier: banner + report + disclosure + rate card + proposals."""
    from docx import Document

    from app.exports.docx import (
        _add_data_table,
        _add_key_value_table,
        _add_subheading,
        _document_bytes,
        build_report_document,
    )

    report = ctx.get("report") or {}

    if report.get("labels"):
        document = build_report_document(report)
    else:
        document = Document()
        _add_fallback_report_docx(document, report)

    banner_text = ctx.get("internal_banner") or INTERNAL_BANNER
    if document.paragraphs:
        banner_paragraph = document.paragraphs[0].insert_paragraph_before(banner_text)
        banner_paragraph.style = document.styles["Heading 1"]
    else:
        document.add_heading(banner_text, level=1)

    document.add_heading("Internal Disclosure Appendix", level=2)
    _add_disclosure_docx(document, report)

    document.add_heading("Rate Card Appendix", level=2)
    rate_card = ctx.get("rate_card")
    _add_key_value_table(document, _rate_card_role_rows(rate_card))
    settings = (rate_card or {}).get("settings") or {}
    role_table = _rate_card_role_table(rate_card)
    if role_table:
        _add_data_table(document, role_table[0], role_table[1])
    setup_table = _rate_card_item_table(
        settings, keys=("setup_cost_items", "nrc_items"), headers=["NRC Item", "Cost"]
    )
    if setup_table:
        _add_data_table(document, setup_table[0], setup_table[1])
    monthly_table = _rate_card_item_table(
        settings, keys=("monthly_rc_items", "rc_items"), headers=["RC Item", "Monthly"]
    )
    if monthly_table:
        _add_data_table(document, monthly_table[0], monthly_table[1])

    document.add_heading("Proposal Appendix", level=2)
    proposals = ctx.get("proposals") or []
    if not proposals:
        document.add_paragraph("No proposal (none)")
    else:
        for proposal in proposals:
            _add_subheading(
                document,
                f"Proposal ({proposal.get('locale', '')}) — {proposal.get('status', '')}",
            )
            _add_proposal_part_docx(document, "Assessment", proposal.get("assessment"))
            _add_proposal_part_docx(document, "Proposal Body", proposal.get("proposal_body"))
            _add_proposal_part_docx(document, "POC", proposal.get("poc"))

    return _document_bytes(document)


def _write_rate_card_sheet(ws: Any, rate_card: dict[str, Any] | None) -> None:
    row_idx = 1
    for label, value in _rate_card_role_rows(rate_card):
        ws.cell(row=row_idx, column=1, value=label)
        ws.cell(row=row_idx, column=2, value=value)
        row_idx += 1
    row_idx += 1

    settings = (rate_card or {}).get("settings") or {}
    role_table = _rate_card_role_table(rate_card)
    if role_table:
        headers, rows = role_table
        for col_idx, header in enumerate(headers, start=1):
            ws.cell(row=row_idx, column=col_idx, value=header)
        row_idx += 1
        for row in rows:
            for col_idx, value in enumerate(row, start=1):
                ws.cell(row=row_idx, column=col_idx, value=value)
            row_idx += 1
        row_idx += 1

    for keys, headers in (
        (("setup_cost_items", "nrc_items"), ["NRC Item", "Cost"]),
        (("monthly_rc_items", "rc_items"), ["RC Item", "Monthly"]),
    ):
        item_table = _rate_card_item_table(settings, keys=keys, headers=headers)
        if not item_table:
            continue
        table_headers, rows = item_table
        for col_idx, header in enumerate(table_headers, start=1):
            ws.cell(row=row_idx, column=col_idx, value=header)
        row_idx += 1
        for row in rows:
            for col_idx, value in enumerate(row, start=1):
                ws.cell(row=row_idx, column=col_idx, value=value)
            row_idx += 1
        row_idx += 1


def _write_disclosure_sheet(ws: Any, report: dict[str, Any]) -> None:
    """XLSX rendering of the Internal Disclosure Appendix: fields client reports omit."""
    extracted = report.get("extracted") or {}
    row_idx = 1

    ws.cell(row=row_idx, column=1, value="Cost Drivers")
    row_idx += 1
    cost_driver_rows = _cost_driver_rows(extracted.get("cost_drivers"))
    if cost_driver_rows:
        ws.cell(row=row_idx, column=1, value="Driver")
        ws.cell(row=row_idx, column=2, value="Impact (JPY)")
        row_idx += 1
        for name, impact in cost_driver_rows:
            ws.cell(row=row_idx, column=1, value=name)
            ws.cell(row=row_idx, column=2, value=impact)
            row_idx += 1
    else:
        ws.cell(row=row_idx, column=1, value="none")
        row_idx += 1
    row_idx += 1

    for key, label in _DISCLOSURE_LIST_FIELDS:
        ws.cell(row=row_idx, column=1, value=label)
        row_idx += 1
        items = extracted.get(key) or []
        if items:
            for item in items:
                ws.cell(row=row_idx, column=1, value=str(item))
                row_idx += 1
        else:
            ws.cell(row=row_idx, column=1, value="none")
            row_idx += 1
        row_idx += 1

    ws.cell(row=row_idx, column=1, value="Confidence")
    row_idx += 1
    score = extracted.get("confidence_score")
    ws.cell(row=row_idx, column=1, value="Score")
    ws.cell(row=row_idx, column=2, value=score if score is not None else "none")
    row_idx += 1
    ws.cell(row=row_idx, column=1, value="Notes")
    ws.cell(row=row_idx, column=2, value=extracted.get("confidence_notes") or "none")
    row_idx += 1
    row_idx += 1

    ws.cell(row=row_idx, column=1, value="Questionnaire Appendix")
    row_idx += 1
    questionnaire_sections = report.get("questionnaire_sections") or []
    if questionnaire_sections:
        for section in questionnaire_sections:
            ws.cell(row=row_idx, column=1, value=section.get("title") or "")
            row_idx += 1
            for field in section.get("fields") or []:
                ws.cell(row=row_idx, column=1, value=field.get("label", ""))
                ws.cell(row=row_idx, column=2, value=str(field.get("value", "")))
                row_idx += 1
            row_idx += 1
    else:
        ws.cell(row=row_idx, column=1, value="none")


def _write_proposal_part_sheet(ws: Any, row_idx: int, title: str, blob: dict[str, Any] | None) -> int:
    """Render an assessment/proposal_body/poc blob as structured rows (not str(dict))."""
    ws.cell(row=row_idx, column=1, value=title)
    row_idx += 1
    if not blob:
        ws.cell(row=row_idx, column=1, value="None")
        return row_idx + 1

    brief_rows = brief_field_rows(blob.get("project_brief"), _BRIEF_LABELS)
    if brief_rows:
        ws.cell(row=row_idx, column=1, value="Project Brief")
        row_idx += 1
        for label, value in brief_rows:
            ws.cell(row=row_idx, column=1, value=label)
            ws.cell(row=row_idx, column=2, value=value)
            row_idx += 1

    rows = section_rows(blob)
    for row in rows:
        if row["title"]:
            ws.cell(row=row_idx, column=1, value=row["title"])
            row_idx += 1
        if row["body"]:
            ws.cell(row=row_idx, column=1, value=str(row["body"]))
            row_idx += 1
        for bullet in row["bullets"]:
            ws.cell(row=row_idx, column=1, value=f"- {bullet}")
            row_idx += 1
        if row.get("rating"):
            ws.cell(row=row_idx, column=1, value=f"Rating: {row['rating']}")
            row_idx += 1

    if not brief_rows and not rows:
        ws.cell(row=row_idx, column=1, value="None")
        row_idx += 1

    return row_idx + 1


def _write_proposals_sheet(ws: Any, proposals: list[dict[str, Any]]) -> None:
    row_idx = 1
    if not proposals:
        ws.cell(row=row_idx, column=1, value="No proposal (none)")
        return

    for proposal in proposals:
        ws.cell(
            row=row_idx,
            column=1,
            value=f"Proposal ({proposal.get('locale', '')}) — {proposal.get('status', '')}",
        )
        row_idx += 1
        row_idx = _write_proposal_part_sheet(ws, row_idx, "Assessment", proposal.get("assessment"))
        row_idx = _write_proposal_part_sheet(ws, row_idx, "Proposal Body", proposal.get("proposal_body"))
        row_idx = _write_proposal_part_sheet(ws, row_idx, "POC", proposal.get("poc"))
        row_idx += 1


def generate_internal_xlsx(ctx: dict[str, Any]) -> bytes:
    """Render the internal XLSX dossier: banner + report + disclosure + rate card + proposals."""
    from io import BytesIO

    from openpyxl import Workbook

    from app.exports.excel import add_report_sheets

    report = ctx.get("report") or {}
    banner_text = ctx.get("internal_banner") or INTERNAL_BANNER

    wb = Workbook()
    wb.remove(wb.active)

    internal_ws = wb.create_sheet("Internal")
    internal_ws.cell(row=1, column=1, value=banner_text)

    if report.get("labels"):
        add_report_sheets(wb, report)
    else:
        report_ws = wb.create_sheet("Report")
        project_summary = report.get("project_summary") or {}
        row_idx = 1
        for label, key in (
            ("Project Name", "project_name"),
            ("Client", "client_name"),
            ("Estimate ID", "estimate_id"),
        ):
            value = project_summary.get(key)
            if value:
                report_ws.cell(row=row_idx, column=1, value=label)
                report_ws.cell(row=row_idx, column=2, value=str(value))
                row_idx += 1
        if row_idx == 1:
            report_ws.cell(row=1, column=1, value=MISSING_CALCULATION_WARNING)

    disclosure_ws = wb.create_sheet("Internal Disclosure")
    _write_disclosure_sheet(disclosure_ws, report)

    rate_card_ws = wb.create_sheet("Rate Card Appendix")
    _write_rate_card_sheet(rate_card_ws, ctx.get("rate_card"))

    proposals_ws = wb.create_sheet("Proposal Appendix")
    _write_proposals_sheet(proposals_ws, ctx.get("proposals") or [])

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
