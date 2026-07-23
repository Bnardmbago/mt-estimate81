"""Generate proposal export bytes (PDF/DOCX/MD/XLSX)."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.exports.markdown import format_currency
from app.exports.theme import PRIMARY, TEXT_ON_PRIMARY
from app.proposals.export_pack_content import (
    brief_field_rows,
    collect_diagrams,
    gantt_timeline_rows,
    iter_pack_parts,
)
from app.proposals.svg_raster import svg_to_png_bytes

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "exports" / "templates"


def _env() -> Environment:
    return Environment(
        loader=FileSystemLoader(TEMPLATE_DIR),
        autoescape=select_autoescape(["html", "xml"]),
        trim_blocks=True,
        lstrip_blocks=True,
    )


def _money(value: Any) -> str:
    if value is None:
        return "—"
    try:
        return format_currency(value)
    except Exception:
        return str(value)


def generate_proposal_pdf(ctx: dict[str, Any]) -> bytes:
    html = _env().get_template("proposal_pack.html.j2").render(ctx=ctx, locale=ctx.get("locale", "en"))
    from weasyprint import HTML

    return HTML(string=html, base_url=str(TEMPLATE_DIR)).write_pdf()


def generate_proposal_markdown(ctx: dict[str, Any]) -> bytes:
    labels = ctx.get("labels") or {}
    lines: list[str] = [
        f"# {ctx.get('project_name', '')}",
        "",
        f"{ctx.get('client_name', '')} — {labels.get('title', 'Proposal')}",
        "",
        f"## {labels.get('toc', 'Table of Contents')}",
        "",
    ]
    for item in ctx.get("toc") or []:
        prefix = "- " if item.get("level") == "1" else "  - "
        lines.append(f"{prefix}{item.get('title')}")
    lines.append("")

    costs = ctx.get("cost_summary") or {}

    lines.extend(
        [
            f"- **{labels.get('one_time')}:** {_money(costs.get('one_time_project_cost_jpy'))}",
            f"- **{labels.get('monthly')}:** {_money(costs.get('monthly_recurring_cost_jpy'))}",
            f"- **{labels.get('first_year')}:** {_money(costs.get('first_year_total_jpy'))}",
            "",
        ]
    )

    def emit_part(title: str, blob: dict[str, Any] | None) -> None:
        if not blob:
            return
        lines.append(f"## {title}")
        lines.append("")
        brief = blob.get("project_brief")
        if brief:
            lines.append(f"### {labels.get('project_brief', 'Project brief')}")
            lines.append("")
            lines.append(f"- **{labels.get('brief_project_name')}:** {brief.get('project_name', '')}")
            lines.append(f"- **{labels.get('brief_description')}:** {brief.get('project_description', '')}")
            lines.append(f"- **{labels.get('brief_business_problem')}:** {brief.get('business_problem', '')}")
            lines.append(f"- **{labels.get('brief_target_users')}:** {brief.get('target_users', '')}")
            lines.append(f"- **{labels.get('brief_technology_stack')}:** {brief.get('technology_stack', '')}")
            lines.append(f"- **{labels.get('brief_constraints')}:** {brief.get('constraints', '')}")
            lines.append("")
        for section in blob.get("sections") or []:
            lines.append(f"### {section.get('title')}")
            lines.append("")
            if section.get("body"):
                lines.append(str(section["body"]))
                lines.append("")
            for bullet in section.get("bullets") or []:
                lines.append(f"- {bullet}")
            if section.get("bullets"):
                lines.append("")
        for table in blob.get("tables") or []:
            lines.append(f"### {table.get('title')}")
            lines.append("")
            headers = table.get("headers") or []
            if headers:
                lines.append("| " + " | ".join(str(h) for h in headers) + " |")
                lines.append("| " + " | ".join("---" for _ in headers) + " |")
            for row in table.get("rows") or []:
                lines.append("| " + " | ".join(str(c) for c in row) + " |")
            lines.append("")
        for diagram in blob.get("diagrams") or []:
            lines.append(f"### {diagram.get('title')}")
            lines.append("")
            lines.append("```mermaid")
            lines.append(str(diagram.get("source") or ""))
            lines.append("```")
            lines.append("")
        if blob.get("milestones"):
            lines.append(f"### {labels.get('poc_milestones', 'Milestones')}")
            lines.append("")
            for m in blob["milestones"]:
                lines.append(f"- {m.get('name')} — {m.get('date') or ''}")
            lines.append("")

    emit_part(labels.get("assessment", "Assessment"), ctx.get("assessment"))
    emit_part(labels.get("proposal", "Proposal"), ctx.get("proposal_body"))

    if ctx.get("milestones"):
        lines.append(f"### {labels.get('milestones', 'Milestones')}")
        lines.append("")
        for m in ctx["milestones"]:
            lines.append(f"- {m.get('name')} — {m.get('date') or ''}")
        lines.append("")

    emit_part(labels.get("poc", "Proof of Concept"), ctx.get("poc"))
    poc = ctx.get("poc") or {}
    official = poc.get("official") or {}
    if official:
        lines.extend(
            [
                f"### {labels.get('official_poc_cost')}",
                "",
                f"- {labels.get('effort_hours')}: {official.get('total_effort_hours')}",
                f"- {labels.get('one_time')}: {_money(official.get('estimated_one_time_cost_jpy'))}",
                f"- {labels.get('suggested_window')}: {poc.get('suggested_validation_window')}",
                "",
            ]
        )

    return ("\n".join(lines)).encode("utf-8")


def _docx_set_run_color(paragraph, hex_color: str) -> None:
    from docx.shared import RGBColor

    try:
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    except (ValueError, IndexError):
        return
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _docx_add_monospace(doc, text: str) -> None:
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _docx_picture_size(png: bytes) -> tuple[Any, Any]:
    """Return (width, height) Inches for a compact diagram embed."""
    from docx.shared import Inches

    max_w = 4.2
    max_h = 5.0
    try:
        from PIL import Image

        image = Image.open(io.BytesIO(png))
        width_px, height_px = image.size
        if width_px <= 0 or height_px <= 0:
            return Inches(max_w), None
        aspect = height_px / width_px
        width_in = max_w
        height_in = width_in * aspect
        if height_in > max_h:
            height_in = max_h
            width_in = height_in / aspect
        return Inches(width_in), Inches(height_in)
    except Exception:
        return Inches(max_w), None


def _docx_embed_png(doc, png: bytes) -> None:
    width, height = _docx_picture_size(png)
    if height is None:
        doc.add_picture(io.BytesIO(png), width=width)
    else:
        doc.add_picture(io.BytesIO(png), width=width, height=height)


def _docx_add_diagram(doc, diagram: dict[str, Any]) -> None:
    """Embed Chromium-rendered diagram PNG when available; otherwise Mermaid source."""
    import base64

    doc.add_heading(str(diagram.get("title") or ""), level=3)
    png: bytes | None = None
    b64 = str(diagram.get("png_base64") or "").strip()
    if b64:
        try:
            png = base64.b64decode(b64)
        except Exception:
            png = None
    if not png:
        svg = str(diagram.get("svg") or "").strip()
        if svg:
            png = svg_to_png_bytes(svg)
    if png:
        _docx_embed_png(doc, png)
        return
    _docx_add_monospace(doc, str(diagram.get("source") or ""))


def _docx_add_table(doc, headers: list[Any], rows: list[list[Any]]) -> None:
    cols = max(len(headers), max((len(r) for r in rows), default=0), 1)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx in range(cols):
        header_cells[idx].text = str(headers[idx]) if idx < len(headers) else ""
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx in range(cols):
            cells[c_idx].text = str(row[c_idx]) if c_idx < len(row) else ""


def _docx_emit_part(doc, title: str, blob: dict[str, Any], labels: dict[str, Any]) -> None:
    doc.add_heading(title, level=1)
    brief_rows = brief_field_rows(blob.get("project_brief"), labels)
    if brief_rows:
        doc.add_heading(str(labels.get("project_brief") or "Project brief"), level=2)
        for label, value in brief_rows:
            doc.add_paragraph(f"{label}: {value}")

    for section in blob.get("sections") or []:
        doc.add_heading(str(section.get("title") or ""), level=2)
        if section.get("body"):
            doc.add_paragraph(str(section["body"]))
        for bullet in section.get("bullets") or []:
            doc.add_paragraph(str(bullet), style="List Bullet")

    tables = blob.get("tables") or []
    if tables:
        doc.add_heading(str(labels.get("poc_tables") or "Tables"), level=2)
        for table in tables:
            doc.add_heading(str(table.get("title") or ""), level=3)
            _docx_add_table(doc, list(table.get("headers") or []), list(table.get("rows") or []))

    diagrams = blob.get("diagrams") or []
    if diagrams:
        doc.add_heading(str(labels.get("poc_diagrams") or "Illustrations"), level=2)
        for diagram in diagrams:
            _docx_add_diagram(doc, diagram)

    if blob.get("milestones"):
        doc.add_heading(str(labels.get("poc_milestones") or "Milestones"), level=2)
        for m in blob["milestones"]:
            doc.add_paragraph(f"{m.get('name')} — {m.get('date') or ''}", style="List Bullet")


def generate_proposal_docx(ctx: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Inches

    labels = ctx.get("labels") or {}
    theme = ctx.get("theme") or {}
    primary_hex = str(theme.get("primary") or PRIMARY)

    doc = Document()
    title = doc.add_heading(str(ctx.get("project_name") or "Proposal"), level=0)
    _docx_set_run_color(title, primary_hex)

    subtitle = doc.add_paragraph(f"{ctx.get('client_name', '')} — {labels.get('title', '')}")
    _docx_set_run_color(subtitle, primary_hex)

    doc.add_heading(str(labels.get("toc") or "Table of Contents"), level=1)
    for item in ctx.get("toc") or []:
        doc.add_paragraph(str(item.get("title") or ""), style="List Number")

    costs = ctx.get("cost_summary") or {}
    doc.add_paragraph(f"{labels.get('one_time')}: {_money(costs.get('one_time_project_cost_jpy'))}")
    doc.add_paragraph(f"{labels.get('monthly')}: {_money(costs.get('monthly_recurring_cost_jpy'))}")
    doc.add_paragraph(f"{labels.get('first_year')}: {_money(costs.get('first_year_total_jpy'))}")

    for _key, part_title, blob in iter_pack_parts(ctx):
        if _key == "poc":
            # PoC is emitted after timeline/milestones to match PDF order loosely;
            # still include full content here when iterating — we'll skip poc in
            # the first pass and emit later. Actually PDF emits assessment,
            # proposal, then gantt, milestones, diagrams, then poc. For DOCX
            # parity of *content*, order can follow PDF.
            continue
        _docx_emit_part(doc, part_title, blob, labels)

    if ctx.get("proposal_body") and ctx.get("gantt_svg"):
        doc.add_heading(str(labels.get("timeline") or "Project timeline"), level=1)
        png = svg_to_png_bytes(str(ctx.get("gantt_svg") or ""))
        if png:
            stream = io.BytesIO(png)
            doc.add_picture(stream, width=Inches(6.5))

    if ctx.get("milestones"):
        doc.add_heading(str(labels.get("milestones") or "Milestones"), level=2)
        for m in ctx["milestones"]:
            doc.add_paragraph(f"{m.get('name')} — {m.get('date') or ''}", style="List Bullet")

    # Top-level proposal diagrams (not inside PoC blob)
    top_diagrams = ctx.get("diagrams") or []
    if top_diagrams:
        for diagram in top_diagrams:
            doc.add_heading(str(diagram.get("title") or ""), level=2)
            import base64

            png: bytes | None = None
            b64 = str(diagram.get("png_base64") or "").strip()
            if b64:
                try:
                    png = base64.b64decode(b64)
                except Exception:
                    png = None
            if not png:
                svg = str(diagram.get("svg") or "").strip()
                if svg:
                    png = svg_to_png_bytes(svg)
            if png:
                _docx_embed_png(doc, png)
            else:
                _docx_add_monospace(doc, str(diagram.get("source") or ""))

    poc = ctx.get("poc")
    if poc:
        poc_title = str(labels.get("poc") or "Proof of Concept")
        _docx_emit_part(doc, poc_title, poc, labels)
        official = poc.get("official") or {}
        if official:
            doc.add_heading(str(labels.get("official_poc_cost") or "Official PoC cost"), level=2)
            doc.add_paragraph(f"{labels.get('effort_hours')}: {official.get('total_effort_hours')}")
            doc.add_paragraph(
                f"{labels.get('one_time')}: {_money(official.get('estimated_one_time_cost_jpy'))}"
            )
            if poc.get("suggested_validation_window"):
                doc.add_paragraph(
                    f"{labels.get('suggested_window')}: {poc.get('suggested_validation_window')}"
                )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _xlsx_header_fill():
    from openpyxl.styles import PatternFill

    return PatternFill(fill_type="solid", fgColor=PRIMARY)


def _xlsx_header_font():
    from openpyxl.styles import Font

    return Font(bold=True, color=TEXT_ON_PRIMARY)


def _xlsx_style_header_row(ws, row: int = 1) -> None:
    fill = _xlsx_header_fill()
    font = _xlsx_header_font()
    for cell in ws[row]:
        cell.fill = fill
        cell.font = font


def _xlsx_append_section_sheet(wb, sheet_name: str, blob: dict[str, Any] | None) -> None:
    if not blob:
        return
    ws = wb.create_sheet(sheet_name[:31])
    ws.append(["Section", "Body", "Bullet", "Rating"])
    _xlsx_style_header_row(ws)
    for section in blob.get("sections") or []:
        title = section.get("title") or ""
        body = section.get("body") or ""
        rating = section.get("rating") or ""
        bullets = section.get("bullets") or []
        if not bullets:
            ws.append([title, body, "", rating])
            continue
        for idx, bullet in enumerate(bullets):
            ws.append(
                [
                    title if idx == 0 else "",
                    body if idx == 0 else "",
                    bullet,
                    rating if idx == 0 else "",
                ]
            )


def generate_proposal_xlsx(ctx: dict[str, Any]) -> bytes:
    from openpyxl import Workbook

    labels = ctx.get("labels") or {}
    wb = Workbook()

    # Summary
    ws = wb.active
    ws.title = "Summary"
    ws.append(["Field", "Value"])
    _xlsx_style_header_row(ws)
    costs = ctx.get("cost_summary") or {}
    poc = ctx.get("poc") or {}
    summary_rows = [
        ("Project", ctx.get("project_name")),
        ("Client", ctx.get("client_name")),
        (labels.get("one_time"), costs.get("one_time_project_cost_jpy")),
        (labels.get("monthly"), costs.get("monthly_recurring_cost_jpy")),
        (labels.get("first_year"), costs.get("first_year_total_jpy")),
        (labels.get("suggested_window"), poc.get("suggested_validation_window")),
    ]
    for field, value in summary_rows:
        ws.append([field, value])

    # Assessment / Proposal / PoC section sheets
    _xlsx_append_section_sheet(wb, "Assessment", ctx.get("assessment"))
    _xlsx_append_section_sheet(wb, "Proposal", ctx.get("proposal_body"))

    if poc:
        poc_ws = wb.create_sheet("PoC")
        poc_ws.append(["Field", "Value"])
        _xlsx_style_header_row(poc_ws)
        for label, value in brief_field_rows(poc.get("project_brief"), labels):
            poc_ws.append([label, value])
        official = poc.get("official") or {}
        if official:
            poc_ws.append([labels.get("effort_hours"), official.get("total_effort_hours")])
            poc_ws.append(
                [labels.get("one_time"), official.get("estimated_one_time_cost_jpy")]
            )
            poc_ws.append(
                [labels.get("suggested_window"), poc.get("suggested_validation_window")]
            )
        poc_ws.append([])
        poc_ws.append(["Section", "Body", "Bullet"])
        for section in poc.get("sections") or []:
            title = section.get("title") or ""
            body = section.get("body") or ""
            bullets = section.get("bullets") or []
            if not bullets:
                poc_ws.append([title, body, ""])
                continue
            for idx, bullet in enumerate(bullets):
                poc_ws.append([title if idx == 0 else "", body if idx == 0 else "", bullet])

        tables = poc.get("tables") or []
        if tables:
            tables_ws = wb.create_sheet("PoC Tables")
            for table in tables:
                tables_ws.append([table.get("title") or ""])
                headers = list(table.get("headers") or [])
                if headers:
                    tables_ws.append(headers)
                    _xlsx_style_header_row(tables_ws, tables_ws.max_row)
                for row in table.get("rows") or []:
                    tables_ws.append(list(row))
                tables_ws.append([])

    diagrams = collect_diagrams(ctx)
    if diagrams:
        diag_ws = wb.create_sheet("Diagrams")
        diag_ws.append(["Title", "Source"])
        _xlsx_style_header_row(diag_ws)
        for diagram in diagrams:
            diag_ws.append([diagram.get("title") or "", diagram.get("source") or ""])

    timeline_rows = gantt_timeline_rows(ctx.get("gantt"))
    timeline = wb.create_sheet("Timeline")
    timeline.append(["Name", "Phase", "Start", "End", "Days", "Hours"])
    _xlsx_style_header_row(timeline)
    for row in timeline_rows:
        timeline.append(
            [
                row.get("name"),
                row.get("phase"),
                row.get("start_date"),
                row.get("end_date"),
                row.get("duration_working_days"),
                row.get("hours"),
            ]
        )
    # Fall back to milestones on Timeline if no gantt tasks
    if not timeline_rows:
        for m in ctx.get("milestones") or []:
            timeline.append([m.get("name"), "", m.get("date"), "", "", ""])

    milestones = wb.create_sheet("Milestones")
    milestones.append(["Name", "Date"])
    _xlsx_style_header_row(milestones)
    for m in ctx.get("milestones") or []:
        milestones.append([m.get("name"), m.get("date")])

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
